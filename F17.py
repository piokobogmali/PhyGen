import sys
from PyQt6.QtWidgets import (
    QApplication, QInputDialog, QMainWindow, QLabel, QWidget, QVBoxLayout,
    QLineEdit, QPushButton, QMessageBox, QFrame, QHBoxLayout,
    QComboBox, QFileDialog, QGridLayout, QGroupBox, QButtonGroup, QRadioButton,
    QDialog, QCheckBox, QTextEdit, QScrollArea, QSpinBox, QDoubleSpinBox,
    QListWidget, QListWidgetItem, QStackedWidget, QTableWidget, QTableWidgetItem,
    QProgressBar, QStatusBar, QHeaderView
)
from PyQt6.QtGui import QPixmap, QPalette, QBrush, QColor, QFont
from PyQt6.QtCore import Qt, QSize, QTimer, QPropertyAnimation, QRect
from PyQt6.QtWidgets import QDialog, QLabel, QVBoxLayout, QProgressBar, QPushButton
import random
import json
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hashlib
from json_database import validate_user, add_user, reset_admin_password
import fitz  # PyMuPDF
import os
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl import Workbook

DEFAULT_QUESTION_CLASSES = ["multiple_choice", "true_false", "problem_solving"]

def load_question_templates(file_path):
    """Load question templates from a JSON file."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return {}

def reset_admin_password(new_password):
    """Reset the admin password in the database."""
    with open("database.json", "r") as file:
        data = json.load(file)
    data["admin_password"] = hashlib.md5(new_password.encode()).hexdigest()
    with open("database.json", "w") as file:
        json.dump(data, file, indent=4)

reset_admin_password("admin123")

# Centralized Styles
BASE_FONT_SIZE = 14  # Adjust for readability

GLOBAL_STYLE = f"""
    QWidget {{
        font-family: 'Segoe UI', Arial, sans-serif;
        font-size: {BASE_FONT_SIZE}px;
        color: #333333;  /* Neutral dark gray for text */
    }}
    QLabel {{
        font-size: {BASE_FONT_SIZE}px;
        color: #333333;
    }}
    QPushButton {{
        background-color: #4CAF50;  /* Green */
        color: white;
        border: 1px solid #4CAF50;
        border-radius: 5px;
        padding: 8px 16px;
        font-weight: bold;
    }}
    QPushButton:hover {{
        background-color: #388E3C;  /* Darker green */
    }}
    QLineEdit, QComboBox, QSpinBox, QDoubleSpinBox {{
        background-color: #fff;
        color: #222;
        border: 1px solid #E0E0E0;
        border-radius: 5px;
        padding: 10px;
        font-size: 16px;
    }}
    QLineEdit:focus, QComboBox:focus, QSpinBox:focus, QDoubleSpinBox:focus {{
        border: 1px solid #4CAF50;  /* Green border on focus */
    }}
    QTableWidget {{
        background-color: #FFFFFF;
        border: 1px solid #E0E0E0;
        border-radius: 5px;
        font-size: {BASE_FONT_SIZE}px;
        color: #333333;
    }}
    QHeaderView::section {{
        background-color: #F5F5F5;
        color: #333333;
        font-weight: bold;
        font-size: {BASE_FONT_SIZE}px;
        border: none;
        padding: 5px;
    }}
    QTableWidget::item {{
        border: none;
        padding: 5px;
    }}
"""

# Button Styles
COMMON_BUTTON_STYLE = """
    QPushButton {{
        background-color: %s;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 14px;
        font-weight: bold;
    }}
    QPushButton:hover {{
        background-color: %s;
    }}
"""

# Button Colors
PRIMARY_BUTTON_COLOR = "#4CAF50"  # Green
PRIMARY_BUTTON_HOVER_COLOR = "#388E3C"  # Darker green
SECONDARY_BUTTON_COLOR = "#2196F3"  # Blue
SECONDARY_BUTTON_HOVER_COLOR = "#1976D2"  # Darker blue
WARNING_BUTTON_COLOR = "#F44336"  # Red
WARNING_BUTTON_HOVER_COLOR = "#D32F2F"  # Darker red

DARK_GREEN = "#008000"

GROUP_BOX_STYLE = f"""
    QGroupBox {{
        font-size: {BASE_FONT_SIZE + 2}px;
        font-weight: bold;
        color: {DARK_GREEN};
        border: 1px solid {DARK_GREEN};
        border-radius: 7px;
        margin-top: 10px;
        padding: 10px;
        background: transparent;
    }}
"""

class AdminDashboard(QWidget):
    def __init__(self, parent=None, user_role="admin"):
        super().__init__(parent)
        self.parent_window = parent
        self.user_role = user_role
        layout = QVBoxLayout(self)
        title_label = QLabel("Admin Dashboard")
        title_label.setStyleSheet(f"font-size: {BASE_FONT_SIZE + 8}px; font-weight: bold; margin-bottom: 15px; color: {DARK_GREEN};")
        layout.addWidget(title_label)

        welcome_label = QLabel("Welcome to PhyGen - Physics Exam Generator")
        welcome_label.setStyleSheet(f"color: {DARK_GREEN}; font-size: 18px;")
        layout.addWidget(welcome_label)

        # Quick actions
        quick_actions = QGroupBox("Quick Actions")
        quick_actions.setStyleSheet(GROUP_BOX_STYLE)
        quick_layout = QHBoxLayout()
        new_exam_btn = QPushButton("New Exam")
        new_exam_btn.setStyleSheet(f"""
            QPushButton {{
                background: #5fd17a;
                color: {DARK_GREEN};
                font-weight: bold;
                border-radius: 7px;
                font-size: 16px;
            }}
            QPushButton:hover {{
                background: #008000;
                color: #fff;
            }}
        """)
        new_exam_btn.clicked.connect(self.open_new_exam)
        quick_layout.addWidget(new_exam_btn)

        # Only show Manage Users for admin
        if self.user_role == "admin":
            manage_users_btn = QPushButton("Manage Users")
            manage_users_btn.setStyleSheet(new_exam_btn.styleSheet())
            manage_users_btn.clicked.connect(self.open_manage_users)
            quick_layout.addWidget(manage_users_btn)

            # --- Passcode Button ---
            passcode_btn = QPushButton("Set/Delete Passcode")
            passcode_btn.setStyleSheet(new_exam_btn.styleSheet())
            passcode_btn.clicked.connect(self.open_passcode_dialog)
            quick_layout.addWidget(passcode_btn)

        quick_actions.setLayout(quick_layout)
        layout.addWidget(quick_actions)

        # Stats
        self.stats_group = QGroupBox("Statistics")
        self.stats_group.setStyleSheet(GROUP_BOX_STYLE)
        self.stats_layout = QGridLayout()
        self.stats_group.setLayout(self.stats_layout)
        layout.addWidget(self.stats_group)

        self.easy_label = self._stat_label("Beginner: 0")
        self.inter_label = self._stat_label("Intermediate: 0")
        self.hard_label = self._stat_label("Advanced: 0")
        self.total_label = self._stat_label("Total Exams: 0")

        self.stats_layout.addWidget(QLabel("Questions:"), 0, 0)
        self.stats_layout.addWidget(self.easy_label, 0, 1)
        self.stats_layout.addWidget(self.inter_label, 0, 2)
        self.stats_layout.addWidget(self.hard_label, 0, 3)
        self.stats_layout.addWidget(self.total_label, 1, 0, 1, 4)

        layout.addStretch()
        self.update_stats()

    def _stat_label(self, text):
        lbl = QLabel(text)
        lbl.setStyleSheet(f"font-size: 16px; color: {DARK_GREEN}; font-weight: normal; background: transparent;")
        return lbl

    def update_stats(self):
        easy = self.count_questions("question_templates_beginner.json")
        inter = self.count_questions("question_templates_intermediate.json")
        hard = self.count_questions("question_templates_advanced.json")
        total_exams = self.count_exams()
        self.easy_label.setText(f"Beginner: {easy}")
        self.inter_label.setText(f"Intermediate: {inter}")
        self.hard_label.setText(f"Advanced: {hard}")
        self.total_label.setText(f"Total Exams: {total_exams}")

    def count_questions(self, file_path):
        if not os.path.exists(file_path):
            return 0
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            seen = set()
            for topic, subtopics in data.items():
                for subtopic, classes in subtopics.items():
                    for qclass, qlist in classes.items():
                        for q in qlist:
                            # Use (topic, subtopic, question text) as unique key
                            key = (topic.strip(), subtopic.strip(), q.get("question", "").strip())
                            seen.add(key)
            return len(seen)
        except Exception:
            return 0

    def count_exams(self):
        if not os.path.exists("exam_bank.json"):
            return 0
        try:
            with open("exam_bank.json", "r", encoding="utf-8") as f:
                exams = json.load(f)
            return len(exams)
        except Exception:
            return 0

    def open_new_exam(self):
        # Switch to the New Exam window if parent has the method
        main_window = self.parent_window or self.parent()
        if main_window and hasattr(main_window, "show_new_exam_window"):
            main_window.show_new_exam_window()

    def open_manage_users(self):
        if self.user_role != "admin":
            QMessageBox.warning(self, "Access Denied", "Sorry, only admin can access this feature.")
            return
        dialog = UserManagementDialog(self)
        dialog.exec()

    def open_passcode_dialog(self):
        passcode = get_admin_passcode()
        if passcode:
            # Require passcode to change/remove
            entered, ok = QInputDialog.getText(self, "Enter Passcode", "Enter current admin passcode:", QLineEdit.EchoMode.Password)
            if not (ok and entered == passcode):
                if ok:
                    QMessageBox.critical(self, "Access Denied", "Incorrect passcode.")
                return
        # Ask to set new or remove
        action, ok = QInputDialog.getItem(
            self, "Passcode Options", "Choose action:",
            ["Set New Passcode", "Remove Passcode"], 0, False
        )
        if not ok:
            return
        if action == "Set New Passcode":
            new_pass, ok = QInputDialog.getText(self, "Set Passcode", "Enter new passcode:", QLineEdit.EchoMode.Password)
            if ok and new_pass:
                set_admin_passcode(new_pass)
                QMessageBox.information(self, "Passcode Set", "Admin passcode has been set.")
        elif action == "Remove Passcode":
            set_admin_passcode(None)
            QMessageBox.information(self, "Passcode Removed", "Admin passcode has been removed.")

class SubTopicSelectionWidget(QWidget):
    def __init__(self, main_topic, sub_topics, parent=None):
        super().__init__(parent)
        self.main_topic = main_topic
        layout = QVBoxLayout(self)
        topic_group = QGroupBox(main_topic)
        topic_group.setStyleSheet(GROUP_BOX_STYLE)
        topic_group_layout = QVBoxLayout(topic_group)
        self.sub_topic_checkboxes = {}
        for sub_topic in sub_topics:
            checkbox = QCheckBox(sub_topic)
            checkbox.setStyleSheet(f"color: white; font-size: {BASE_FONT_SIZE}px;")
            topic_group_layout.addWidget(checkbox)
            self.sub_topic_checkboxes[sub_topic] = checkbox
        layout.addWidget(topic_group)

    def get_selected_sub_topics(self):
        return [sub_topic for sub_topic, checkbox in self.sub_topic_checkboxes.items() if checkbox.isChecked()]

from PyQt6.QtWidgets import QAbstractSpinBox

class ExamTypeWidget(QWidget):
    def __init__(self, type_name, description, parent=None):
        super().__init__(parent)
        self.type_name = type_name
        layout = QHBoxLayout(self)
        self.checkbox = QCheckBox(type_name)
        self.checkbox.setStyleSheet(f"font-weight: bold; color: white; font-size: {BASE_FONT_SIZE}px;")
        layout.addWidget(self.checkbox)
        self.percentage_spinbox = QDoubleSpinBox()
        self.percentage_spinbox.setRange(0, 100)
        self.percentage_spinbox.setValue(0)
        self.percentage_spinbox.setSuffix("%")
        self.percentage_spinbox.setEnabled(False)
        self.percentage_spinbox.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.percentage_spinbox.setStyleSheet(f"""
            QDoubleSpinBox {{
                color: white;
                font-size: {BASE_FONT_SIZE}px;
                min-width: 55px;
                max-width: 55px;
                background: rgba(255,255,255,0.08);
                border: 1px solid #BDBDBD;
                border-radius: 5px;
                padding-right: 5px;
            }}
            QDoubleSpinBox::up-button, QDoubleSpinBox::down-button {{
                width: 0px;
                height: 0px;
                border: none;
            }}
        """)
        layout.addWidget(self.percentage_spinbox)
        desc_label = QLabel(description)
        desc_label.setStyleSheet(f"color: #BDBDBD; font-size: {BASE_FONT_SIZE}px;")
        layout.addWidget(desc_label)
        layout.addStretch()
        self.checkbox.stateChanged.connect(self.toggle_percentage_spinbox)

    def toggle_percentage_spinbox(self, state):
        self.percentage_spinbox.setEnabled(state == Qt.CheckState.Checked.value)
        if state != Qt.CheckState.Checked.value:
            self.percentage_spinbox.setValue(0)

class ExamCard(QFrame):
    def __init__(self, exam_data=None, parent=None):
        super().__init__(parent)
        self.exam_data = exam_data
        self.setup_ui()
        
    def setup_ui(self):
        self.setStyleSheet("""
            QFrame {
                background-color: #fff;
                border: 1px solid #E0E0E0;
                border-radius: 8px;
            }
            QFrame:hover {
                background-color: rgba(255, 255, 255, 0.15);
                border: 1px solid rgba(255, 255, 255, 0.3);
            }
        """)
        self.setFixedSize(250, 150)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        
        if self.exam_data:
            # Exam title
            title_label = QLabel(f"Slot No: {self.exam_data.get('id', 'N/A')}")
            title_label.setStyleSheet(f"""
                QLabel {{
                    font-size: {BASE_FONT_SIZE + 2}px;
                    font-weight: bold;
                    color: white;
                }}
            """)
            title_label.setWordWrap(True)
            layout.addWidget(title_label)
            
            # Exam details
            details_label = QLabel(
                f"Date Created: {self.exam_data.get('date', 'N/A')}\n"
                f"Questions: {self.exam_data.get('question_count', 0)}\n"
                f"Type: {self.exam_data.get('type', 'Mixed')}"
            )
            details_label.setStyleSheet(f"""
                QLabel {{
                    font-size: {BASE_FONT_SIZE}px;
                    color: #BDBDBD;
                }}
            """)
            layout.addWidget(details_label)
            
            # Action buttons
            button_layout = QHBoxLayout()
            
            view_btn = QPushButton("View")
            view_btn.setStyleSheet(COMMON_BUTTON_STYLE % ("rgba(255, 255, 255, 0.2)", "rgba(255, 255, 255, 0.3)"))
            view_btn.setFixedHeight(25)
            button_layout.addWidget(view_btn)
            
            menu_btn = QPushButton("?")
            menu_btn.setStyleSheet(COMMON_BUTTON_STYLE % ("rgba(255, 255, 255, 0.2)", "rgba(255, 255, 255, 0.3)"))
            menu_btn.setFixedSize(25, 25)
            button_layout.addWidget(menu_btn)
            
            layout.addLayout(button_layout)
        else:
            # Empty slot with plus sign
            plus_label = QLabel("+")
            plus_label.setStyleSheet(f"""
                QLabel {{
                    font-size: {BASE_FONT_SIZE + 24}px;
                    font-weight: bold;
                    color: rgba(255, 255, 255, 0.5);
                    qproperty-alignment: AlignCenter;
                }}
            """)
            layout.addWidget(plus_label)
            layout.addStretch()

class ExamBankWindow(QMainWindow):
    EXAMS_FILE = "exam_bank.json"

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Exam Bank")
        self.setMinimumSize(900, 700)  # Enlarged window

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        title_label = QLabel("Exam Bank")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 20px;
                font-weight: normal;
                color: #222;
                text-align: center;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        # --- Button Row (Import & Delete) ---
        button_row = QHBoxLayout()
        button_row.addStretch()

        download_btn = QPushButton("Import as Word File")
        download_btn.setStyleSheet("""
QPushButton {
    background-color: #fff;
    color: #222;
    border: 1px solid #E
    border-radius: 5px;
    padding: 10px;
    font-size: 14px;
    font-weight: normal;
}
QPushButton:hover {
    background-color: #f5f5f5;
}
""")
        download_btn.clicked.connect(self.download_exam)
        button_row.addWidget(download_btn)

        # --- Delete Button ---
        delete_btn = QPushButton("Delete Exam")
        delete_btn.setStyleSheet(download_btn.styleSheet())
        delete_btn.clicked.connect(self.delete_exam)
        button_row.addWidget(delete_btn)

        layout.addLayout(button_row)

        self.exam_slots_widget = QWidget()
        self.exam_slots_layout = QGridLayout(self.exam_slots_widget)
        self.exam_slots_layout.setSpacing(25)  # More space between boxes
        self.exam_slots_layout.setContentsMargins(10, 10, 10, 10)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(self.exam_slots_widget)
        scroll.setStyleSheet("QScrollArea {background: #2c2c2c; border: none;}")
        layout.addWidget(scroll)

        nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton("Previous")
        self.next_btn = QPushButton("Next")
        self.prev_btn.setStyleSheet("""
            QPushButton {
                background-color: #009944;
                color: white;
                border-radius: 10px;
                font-weight: bold;
                font-size: 28px;
                padding: 18px;
            }
            QPushButton:hover {
                background-color: #006633;
                color: white;
            }
        """)
        self.next_btn.setStyleSheet("""
    QPushButton {
        background: #fff;
        color: #222;
        border: 2px solid #009944;
        border-radius: 8px;
        font-size: 22px;
        font-weight: normal;
        padding: 10px 0;
    }
    QPushButton:hover {
        background: #f5f5f5;
        color: #009944;
        border: 2px solid #009944;
    }
""")
        self.prev_btn.setMinimumWidth(200)  # Elongated width
        self.prev_btn.setMinimumHeight(60)  # Elongated height
        self.next_btn.setMinimumWidth(200)  # Elongated width
        self.next_btn.setMinimumHeight(60)  # Elongated height
        self.prev_btn.clicked.connect(self.prev_page)
        self.next_btn.clicked.connect(self.next_page)
        nav_layout.addStretch()
        nav_layout.addWidget(self.prev_btn)
        nav_layout.addWidget(self.next_btn)
        nav_layout.addStretch()
        layout.addLayout(nav_layout)

        self.exams = self.load_exams()
        self.page = 0
        self.boxes_per_page = 9  # Still 9 slots per page
        self.selected_exam_idx = None  # Track selected exam
        self.update_exam_slots()

    def save_exams(self):
        with open(self.EXAMS_FILE, "w", encoding="utf-8") as f:
            json.dump(self.exams, f, indent=2)

    def load_exams(self):
        if os.path.exists(self.EXAMS_FILE):
            with open(self.EXAMS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        return []

    def update_exam_slots(self):
        # Clear layout
        for i in reversed(range(self.exam_slots_layout.count())):
            widget = self.exam_slots_layout.itemAt(i).widget()
            if widget:
                widget.setParent(None)

        # Pagination logic
        start = self.page * self.boxes_per_page
        end = start + self.boxes_per_page
        exams_to_show = self.exams[start:end]

        self.exam_boxes = []  # Track boxes for selection

        # Always show 9 boxes (fill with empty if needed)
        for idx in range(self.boxes_per_page):
            if idx < len(exams_to_show):
                exam = exams_to_show[idx]
                box = self.create_exam_box(exam, start + idx)
            else:
                box = self.create_exam_box(None, None)
            row = idx // 3
            col = idx % 3
            self.exam_slots_layout.addWidget(box, row, col)
            self.exam_boxes.append(box)

        self.prev_btn.setEnabled(self.page > 0)
        self.next_btn.setEnabled(end < len(self.exams))

    def create_exam_box(self, exam, global_idx):
        box = QFrame()
        box.setFixedSize(320, 180)  # Enlarged box size
        box.setStyleSheet("""
            QFrame {
                background-color: #fff;
                border: 1px solid #E0E0E0;
                border-radius: 8px;
            }
        """)
        vbox = QVBoxLayout(box)
        vbox.setContentsMargins(16, 16, 16, 16)
        vbox.setSpacing(8)

        if exam:
            vbox.addWidget(self._label(f"Name: {exam['title']}"))
            vbox.addWidget(self._label(f"Date of Creation: {exam['date']}"))
            vbox.addWidget(self._label(f"Type of Exam: {exam['type']}"))
            vbox.addWidget(self._label(f"Difficulty: {exam['difficulty']}"))

            # Make box clickable for selection and answer key
            def on_box_click(event, idx=global_idx):
                # Show answer key on double click
                if event.type() == 2:  # QEvent.MouseButtonDblClick
                    self.show_answer_key(exam)
            box.mousePressEvent = lambda event: on_box_click(event)
        else:
            vbox.addWidget(self._label("Name:"))
            vbox.addWidget(self._label("Date of Creation:"))
            vbox.addWidget(self._label("Type of Exam:"))
            vbox.addWidget(self._label("Difficulty:"))

        return box

    def show_answer_key(self, exam):
        """Display the answer key for the selected exam."""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Answer Key - {exam['title']}")
        dialog.setFixedSize(600, 400)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)

        title_label = QLabel(f"Answer Key for {exam['title']}")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: black;
                text-align: center;
            }
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        answer_list = QListWidget()
        answer_list.setStyleSheet("""
            QListWidget {
                background-color: #fff;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 5px;
                font-size: 14px;
                color: black;
            }
        """)
        for idx, question in enumerate(exam["questions"], start=1):
            answer_list.addItem(f"{idx}. {question['question']} - Answer: {question['answer']}")
        layout.addWidget(answer_list)

        ok_button = QPushButton("OK")
        ok_button.setStyleSheet("""
            QPushButton {
                background-color: #fff;
                color: black;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f5f5f5;
            }
        """)
        ok_button.clicked.connect(dialog.accept)
        layout.addWidget(ok_button, alignment=Qt.AlignmentFlag.AlignCenter)

        dialog.exec()

    def _label(self, text):
        """Helper method to create a QLabel with consistent styling."""
        lbl = QLabel(text)
        lbl.setStyleSheet("color: #222; font-size: 15px; font-weight: normal;")
        return lbl

    def next_page(self):
        if (self.page + 1) * self.boxes_per_page < len(self.exams):
            self.page += 1
            self.selected_exam_idx = None
            self.update_exam_slots()

    def prev_page(self):
        if self.page > 0:
            self.page -= 1
            self.selected_exam_idx = None
            self.update_exam_slots()

    def download_exam(self):
        if not self.exams:
            QMessageBox.information(self, "No Exams", "No exams to download.")
            return

        # Let user pick a slot
        items = [f"{i+1}: {exam['title']}" for i, exam in enumerate(self.exams)]
        item, ok = QInputDialog.getItem(self, "Download Exam", "Select exam slot:", items, 0, False)
        if ok and item:
            idx = int(item.split(":")[0]) - 1
            exam = self.exams[idx]
            path, _ = QFileDialog.getSaveFileName(self, "Save Exam", f"{exam['title']}.docx", "Word Files (*.docx)")
            if path:
                # Show the transferring animation
                dialog = GenerateExamDialog(self, message="Transferring...")
                dialog.start_animation()
                dialog.exec()

                # Export the exam to a Word file
                export_exam_to_word(exam['title'], exam['type'], exam['questions'], path)

                # Export the answer key to a separate Word file
                answer_key_path = path.replace(".docx", "_answer_key.docx")
                export_answer_key_to_word(exam['title'], exam['questions'], answer_key_path)

                # --- Export TOS ---
                tos_path, _ = QFileDialog.getSaveFileName(self, "Save TOS", f"{exam['title']}_TOS.xlsx", "Excel Files (*.xlsx)")
                if tos_path:
                    # Extract topics and subtopics from questions
                    pairs = exam.get('selected_topic_subtopic_pairs', [])
                    if not pairs:
                        # fallback for old exams
                        topics = exam.get('selected_topics', [])
                        subtopics = exam.get('selected_subtopics', [])
                        pairs = list(zip(topics, subtopics))
                                
                    # Count question types for distribution
                    qtype_map = {
                        "Remembering": 0,
                        "Inferential": 0,
                        "Applied": 0
                    }
                    for q in exam['questions']:
                        # You may need to adjust this mapping based on your actual question type names
                        qtype = q.get('type', '').lower()
                        if "remember" in qtype:
                            qtype_map["Remembering"] += 1
                        elif "infer" in qtype or "understand" in qtype or "analyz" in qtype:
                            qtype_map["Inferential"] += 1
                        elif "appl" in qtype or "evaluat" in qtype or "creat" in qtype or "problem" in qtype:
                            qtype_map["Applied"] += 1

                    export_tos_to_excel(
                        tos_path,
                        [t for t, s in pairs],
                        [s for t, s in pairs],
                        len(exam['questions']),
                        exam.get('type', ''),
                        exam.get('difficulty', ''),
                        qtype_map,
                        questions=exam['questions']
                    )
                    QMessageBox.information(self, "Downloaded", f"Exam saved as {path}\nAnswer key saved as {answer_key_path}\nTOS saved as {tos_path}")
                else:
                    QMessageBox.information(self, "Downloaded", f"Exam saved as {path}\nAnswer key saved as {answer_key_path}\nTOS not saved (cancelled).")

        self.save_exams()
        self.update_exam_slots()

    def delete_exam(self):
        if get_admin_passcode() and not require_admin_passcode(self):
            return
        if not self.exams:
            QMessageBox.information(self, "No Exams", "No exams to delete.")
            return

        # Let user pick a slot
        items = [f"{i+1}: {exam['title']}" for i, exam in enumerate(self.exams)]
        item, ok = QInputDialog.getItem(self, "Delete Exam", "Select exam slot to delete:", items, 0, False)
        if ok and item:
            idx = int(item.split(":")[0]) - 1
            exam = self.exams[idx]
            confirm = QMessageBox.question(
                self,
                "Confirm Delete",
                f"Are you sure you want to delete '{exam['title']}'?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if confirm == QMessageBox.StandardButton.Yes:
                del self.exams[idx]
                self.save_exams()
                self.selected_exam_idx = None
                self.update_exam_slots()
                QMessageBox.information(self, "Deleted", f"Exam '{exam['title']}' deleted.")

class NewExamWindow(QMainWindow):
    def __init__(self, exam_bank_window):
        super().__init__()
        self.setWindowTitle("Create New Exam")
        self.setMinimumSize(900, 700)
        self.exam_bank_window = exam_bank_window

        # Main Layout
        main_layout = QVBoxLayout()
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)

        # Swiping/Arrow Navigation
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(self.stacked_widget)

        # --- 1. Basic Information ---
        self.basic_info_widget = QWidget()
        basic_info_layout = QVBoxLayout(self.basic_info_widget)
        basic_info_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        basic_info_group = QGroupBox("Basic Information")
        basic_info_group.setFixedSize(820, 380)  # Rectangle: wider than tall
        basic_info_group.setStyleSheet("""
    QGroupBox {
        font-size: 38px;
        font-weight: bold;
        color: #009944;
        border: 3px solid #009944;
        border-radius: 24px;
        padding: 32px 48px;
        background-color: #fff;
        margin-top: 18px;
    }
""")

        basic_info_form = QGridLayout(basic_info_group)
        basic_info_form.setHorizontalSpacing(50)
        basic_info_form.setVerticalSpacing(40)

        # Exam Name
        self.exam_name_edit = QLineEdit()
        self.exam_name_edit.setPlaceholderText("Enter Exam Name")
        self.exam_name_edit.setMinimumHeight(70)
        self.exam_name_edit.setStyleSheet("""
    QLineEdit {
        background-color: #fff;
        color: #222;
        border: 2px solid #E0E0E0;
        border-radius: 12px;
        padding: 22px;
        font-size: 32px;
    }
    QLineEdit:focus {
        border: 2px solid #4CAF50;
    }
""")
        exam_name_label = QLabel("Exam Name:")
        exam_name_label.setStyleSheet("font-size: 32px; color: #222; font-weight: bold;")
        basic_info_form.addWidget(exam_name_label, 0, 0)
        basic_info_form.addWidget(self.exam_name_edit, 0, 1)

        # Exam Type
        self.exam_type_combobox = QComboBox()
        self.exam_type_combobox.addItems(["Midterm", "Final", "Quiz"])
        self.exam_type_combobox.setMinimumHeight(70)
        self.exam_type_combobox.setStyleSheet("""
    QComboBox {
        background-color: #fff;
        color: #222;
        border: 2px solid #E0E0E0;
        border-radius: 12px;
        padding: 22px;
        font-size: 32px;
        min-width: 220px;
    }
    QComboBox:focus {
        border: 2px solid #4CAF50;
    }
""")
        exam_type_label = QLabel("Type:")
        exam_type_label.setStyleSheet("font-size: 32px; color: #222; font-weight: bold;")
        basic_info_form.addWidget(exam_type_label, 1, 0)
        basic_info_form.addWidget(self.exam_type_combobox, 1, 1)

        basic_info_group.setLayout(basic_info_form)
        basic_info_layout.addStretch()
        basic_info_layout.addWidget(basic_info_group, alignment=Qt.AlignmentFlag.AlignCenter)
        basic_info_layout.addStretch()
        self.stacked_widget.addWidget(self.basic_info_widget)

        # --- 2. Topics and Difficulty with Subtopics ---
        self.topics_widget = QWidget()
        topics_layout = QVBoxLayout(self.topics_widget)
        topics_group = QGroupBox("Topics and Difficulty")
        topics_group.setStyleSheet("""
            QGroupBox {
                font-size: 16px;
                font-weight: normal;
                color: #222;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
            }
        """)
        topics_form = QGridLayout(topics_group)

        # Main topics and their subtopics
        self.topic_subtopics = {
            "Mechanics": [
                "Kinematics", "Dynamics", "Work, Energy, Power", "Momentum", "Circular Motion", "Gravitation"
            ],
            "Waves and Acoustics": [
                "Wave Properties", "Sound", "Doppler Effect"
            ],
            "Thermodynamics": [
                "Temperature", "Heat Transfer", "Laws of Thermodynamics"
            ],
            "Electricity and Magnetism": [
                "Electrostatics", "Current Electricity", "Magnetism", "Electromagnetic Induction"
            ],
            "Optics": [
                "Reflection", "Refraction", "Lenses", "Mirrors"
            ],
            "Modern Physics": [
                "Relativity", "Quantum Physics", "Nuclear Physics"
            ]
        }

        # Topic list
        self.topics_list = QListWidget()
        self.topics_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        self.topics_list.addItems(list(self.topic_subtopics.keys()))
        self.topics_list.setMinimumHeight(220)
        self.topics_list.setStyleSheet("""
            QListWidget {
                background-color: #fff;  /* White background */
                color: #000;  /* Black font */
                border: 1px solid #ccc;  /* Thin gray border */
                border-radius: 5px;
                font-size: 16px;
                padding: 8px;
            }
            QListWidget::item {
                padding: 5px;
            }
            QListWidget::item:hover {
                background-color: #f5f5f5;  /* Light gray background on hover */
            }
""")
        topics_label = QLabel("Select Topic:")
        topics_label.setStyleSheet("font-size: 22px; color: #009944; font-weight: bold;")
        topics_form.addWidget(topics_label, 0, 0)
        topics_form.addWidget(self.topics_list, 0, 1)

        # --- Subtopics area (persistent checkboxes, only visible for selected topic) ---
        self.subtopics_group = QGroupBox("Subtopics")
        self.subtopics_group.setStyleSheet("""
            QGroupBox {
                font-size: 20px;
                font-weight: bold;
                color: #009944;
                border: 2px solid #009944;
                border-radius: 10px;
                margin-top: 10px;
                padding: 10px;
            }
        """)
        self.subtopics_group.setFixedHeight(220)  # <-- Add this line to fix the height (adjust as needed)
        self.subtopics_layout = QVBoxLayout(self.subtopics_group)
        self.subtopic_checkboxes = {}  # {(topic, subtopic): QCheckBox}

        # Create all checkboxes once and keep them persistent
        for topic, subtopics in self.topic_subtopics.items():
            for sub in subtopics:
                key = (topic, sub)
                cb = QCheckBox(sub)
                cb.setStyleSheet("color: #222; font-size: 14px;")
                cb.hide()  # Hide by default
                self.subtopics_layout.addWidget(cb)
                self.subtopic_checkboxes[key] = cb

        def update_subtopics():
            # Hide all checkboxes first
            for cb in self.subtopic_checkboxes.values():
                cb.hide()
            # Show only those for the currently selected topic (single selection)
            selected_items = self.topics_list.selectedItems()
            if selected_items:
                topic = selected_items[0].text()
                for sub in self.topic_subtopics[topic]:
                    self.subtopic_checkboxes[(topic, sub)].show()

        self.topics_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        self.topics_list.itemSelectionChanged.connect(update_subtopics)
        update_subtopics()  # Initialize

        topics_form.addWidget(self.subtopics_group, 1, 0, 1, 2)

        self.difficulty_combobox = QComboBox()
        self.difficulty_combobox.addItems(["Beginner", "Intermediate", "Advanced"])
        self.difficulty_combobox.setMinimumHeight(48)
        self.difficulty_combobox.setStyleSheet("""
    QComboBox {
        background-color: #fff;  /* White background */
        color: #000;  /* Black font */
        border: 1px solid #ccc;  /* Thin gray border */
        border-radius: 5px;
        padding: 10px;
        font-size: 16px;
    }
    QComboBox::drop-down {
        border: none;
    }
    QComboBox:hover {
        background-color: #f5f5f5;  /* Light gray background on hover */
    }
""")
        difficulty_label = QLabel("Difficulty Level:")
        difficulty_label.setStyleSheet("font-size: 14px; color: #222; font-weight: normal;")
        topics_form.addWidget(difficulty_label, 2, 0)
        topics_form.addWidget(self.difficulty_combobox, 2, 1)

        self.num_questions_spinbox = QSpinBox()
        self.num_questions_spinbox.setRange(10, 100)
        self.num_questions_spinbox.setValue(10)
        self.num_questions_spinbox.setMinimumHeight(48)
        self.num_questions_spinbox.setStyleSheet("""
    QSpinBox {
        background-color: #fff;
        color: #222;
        border: 1px solid #E0E0E0;
        border-radius: 5px;
        padding: 10px;
        font-size: 14px;
    }
    QSpinBox:focus {
        border: 1px solid #4CAF50;  /* Green border on focus */
    }
""")
        num_questions_label = QLabel("Number of Questions:")
        num_questions_label.setStyleSheet("font-size: 14px; color: #222; font-weight: normal;")
        topics_form.addWidget(num_questions_label, 3, 0)
        topics_form.addWidget(self.num_questions_spinbox, 3, 1)

        topics_group.setLayout(topics_form)
        topics_layout.addWidget(topics_group)
        topics_layout.addStretch()
        self.stacked_widget.addWidget(self.topics_widget)

        # --- 3. Question Type Distribution + Generate ---
        self.qtype_widget = QWidget()
        qtype_layout = QVBoxLayout(self.qtype_widget)
        question_type_group = QGroupBox("Question Type Distribution")
        question_type_group.setStyleSheet("""
            QGroupBox {
                font-size: 16px;
                font-weight: normal;
                color: #222;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
            }
        """)
        question_type_form = QGridLayout(question_type_group)

        from PyQt6.QtWidgets import QDoubleSpinBox, QAbstractSpinBox

        self.question_types = {
            "Multiple Choice": QDoubleSpinBox(),
            "True/False": QDoubleSpinBox(),
            "Problem Solving": QDoubleSpinBox()
        }
        for row, (qtype, spinbox) in enumerate(self.question_types.items()):
            spinbox.setRange(0, 100)
            spinbox.setDecimals(0)
            spinbox.setSuffix("%")
            spinbox.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
            spinbox.setMinimumHeight(48)
            spinbox.setStyleSheet("""
                QDoubleSpinBox {
                    background-color: #fff;  /* White background */
                    color: #000;  /* Black font */
                    border: 1px solid #ccc;  /* Thin gray border */
                    border-radius: 5px;
                    padding: 10px;
                    font-size: 16px;
                    min-width: 120px;
                    max-width: 160px;
                }
                QDoubleSpinBox:hover {
                    background-color: #f5f5f5;  /* Light gray background on hover */
                }
                QDoubleSpinBox::up-button, QDoubleSpinBox::down-button {
                    width: 0px;  /* Remove up/down buttons */
                    height: 0px;
                    border: none;
                }
            """)
            qtype_label = QLabel(f"{qtype}:")
            qtype_label.setStyleSheet("font-size: 22px; color: #009944; font-weight: bold;")
            question_type_form.addWidget(qtype_label, row, 0)
            question_type_form.addWidget(spinbox, row, 1)

        self.total_percentage_label = QLabel("Total: 0% (must equal 100%)")
        self.total_percentage_label.setStyleSheet("color: #222; font-size: 14px; font-weight: normal;")
        question_type_form.addWidget(self.total_percentage_label, len(self.question_types), 0, 1, 2)

        question_type_group.setLayout(question_type_form)
        qtype_layout.addWidget(question_type_group)

        # --- Cognitive Level Distribution (Remembering, Inferential, Applied) ---
        cognitive_group = QGroupBox("Cognitive Level Distribution")
        cognitive_group.setStyleSheet("""
            QGroupBox {
                font-size: 16px;
                font-weight: normal;
                color: #222;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
            }
        """)
        cognitive_form = QGridLayout(cognitive_group)

        self.cognitive_types = {
            "Remembering": QDoubleSpinBox(),
            "Inferential": QDoubleSpinBox(),
            "Applied": QDoubleSpinBox()
        }
        for row, (ctype, spinbox) in enumerate(self.cognitive_types.items()):
            spinbox.setRange(0, 100)
            spinbox.setDecimals(0)
            spinbox.setSuffix("%")
            spinbox.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
            spinbox.setMinimumHeight(40)
            spinbox.setStyleSheet("""
                QDoubleSpinBox {
                    background-color: #fff;
                    color: #000;
                    border: 1px solid #ccc;
                    border-radius: 5px;
                    padding: 10px;
                    font-size: 16px;
                    min-width: 120px;
                    max-width: 160px;
                }
                QDoubleSpinBox:hover {
                    background-color: #f5f5f5;
                }
                QDoubleSpinBox::up-button, QDoubleSpinBox::down-button {
                    width: 0px;
                    height: 0px;
                    border: none;
                }
            """)
            ctype_label = QLabel(f"{ctype}:")
            ctype_label.setStyleSheet("font-size: 18px; color: #009944; font-weight: bold;")
            cognitive_form.addWidget(ctype_label, row, 0)
            cognitive_form.addWidget(spinbox, row, 1)

        self.cognitive_total_label = QLabel("Total: 0% (must equal 100%)")
        self.cognitive_total_label.setStyleSheet("color: #222; font-size: 14px; font-weight: normal;")
        cognitive_form.addWidget(self.cognitive_total_label, len(self.cognitive_types), 0, 1, 2)

        cognitive_group.setLayout(cognitive_form)
        qtype_layout.addWidget(cognitive_group)

        # Generate Exam Button
        generate_button = QPushButton("Generate Exam")
        generate_button.setMinimumHeight(60)
        generate_button.setStyleSheet("""
            QPushButton {
                background-color: #009944;  /* Dark green */
                color: white;
                border: none;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #006633;  /* Darker green on hover */
            }
        """)
        generate_button.clicked.connect(self.generate_exam)
        qtype_layout.addWidget(generate_button, 0, Qt.AlignmentFlag.AlignCenter)
        qtype_layout.addStretch()
        self.stacked_widget.addWidget(self.qtype_widget)

        # --- Navigation Buttons ---
        nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton("Previous")
        self.next_btn = QPushButton("Next")
        self.prev_btn.setMinimumHeight(60)
        self.next_btn.setMinimumHeight(60)
        self.prev_btn.setMinimumWidth(200)
        self.next_btn.setMinimumWidth(200)
        self.prev_btn.setStyleSheet("""
            QPushButton {
                background-color: #009944;
                color: white;
                border-radius: 10px;
                font-weight: bold;
                font-size: 28px;
                padding: 18px;
            }
            QPushButton:hover {
                background-color: #006633;
                color: white;
            }
        """)
        self.next_btn.setStyleSheet("""
    QPushButton {
        background: #fff;
        color: #222;
        border: 2px solid #009944;
        border-radius: 8px;
        font-size: 22px;
        font-weight: normal;
        padding: 10px 0;
    }
    QPushButton:hover {
        background: #f5f5f5;
        color: #009944;
        border: 2px solid #009944;
    }
""")
        self.prev_btn.setMinimumWidth(200)  # Elongated width
        self.prev_btn.setMinimumHeight(60)  # Elongated height
        self.next_btn.setMinimumWidth(200)  # Elongated width
        self.next_btn.setMinimumHeight(60)  # Elongated height
        self.prev_btn.clicked.connect(self.prev_screen)
        self.next_btn.clicked.connect(self.next_screen)
        nav_layout.addStretch()
        nav_layout.addWidget(self.prev_btn)
        nav_layout.addWidget(self.next_btn)
        nav_layout.addStretch()
        main_layout.addLayout(nav_layout)

        self.stacked_widget.setCurrentIndex(0)
        self.update_nav_buttons()

        # Update total percentage label when spinboxes change
        for spinbox in self.question_types.values():
            spinbox.valueChanged.connect(self.update_total_percentage)
        for spinbox in self.cognitive_types.values():
            spinbox.valueChanged.connect(self.update_cognitive_total_percentage)

    def update_nav_buttons(self):
        idx = self.stacked_widget.currentIndex()
        self.next_btn.setEnabled(idx < self.stacked_widget.count() - 1)

    def prev_screen(self):
        idx = self.stacked_widget.currentIndex()
        if idx > 0:
            self.stacked_widget.setCurrentIndex(idx - 1)
        self.update_nav_buttons()

    def next_screen(self):
        idx = self.stacked_widget.currentIndex()
        if idx == 0:
            # Validate basic info
            if not self.exam_name_edit.text().strip():
                QMessageBox.warning(self, "Validation", "Please enter an exam name.")
                return
        if idx == 1:
            # Validate topics/difficulty
            if not self.topics_list.selectedItems():
                QMessageBox.warning(self, "Validation", "Please select a topic.")
                return
            if not any(cb.isChecked() for cb in self.subtopic_checkboxes.values()):
                QMessageBox.warning(self, "Validation", "Please select at least one subtopic.")
                return

            # --- Show preview dialog before proceeding ---
            checked_subtopics = [((topic, sub), cb)
                                 for (topic, sub), cb in self.subtopic_checkboxes.items() if cb.isChecked()]
            preview = SubtopicPreviewDialog(checked_subtopics, parent=self)
            result = preview.exec()
            if result != QDialog.DialogCode.Accepted:
                return  # User closed the dialog, don't proceed

        if idx < self.stacked_widget.count() - 1:
            self.stacked_widget.setCurrentIndex(idx + 1)
        self.update_nav_buttons()

    def update_total_percentage(self):
        total = sum(spinbox.value() for spinbox in self.question_types.values())
        self.total_percentage_label.setText(f"Total: {int(total)}% (must equal 100%)")
    
    def update_cognitive_total_percentage(self):
        total = sum(spinbox.value() for spinbox in self.cognitive_types.values())
        self.cognitive_total_label.setText(f"Total: {int(total)}% (must equal 100%)")

    def reset_form(self):
        self.exam_name_edit.clear()
        self.exam_type_combobox.setCurrentIndex(0)
        self.topics_list.clearSelection()
        for cb in self.subtopic_checkboxes.values():
            cb.setChecked(False)
        self.difficulty_combobox.setCurrentIndex(0)
        self.num_questions_spinbox.setValue(10)
        for spinbox in self.question_types.values():
            spinbox.setValue(0)
        for spinbox in self.cognitive_types.values():
            spinbox.setValue(0)
        self.stacked_widget.setCurrentIndex(0)
        self.update_nav_buttons()

    def generate_exam(self):
        total_percentage = sum(spinbox.value() for spinbox in self.question_types.values())
        if total_percentage != 100:
            QMessageBox.warning(self, "Invalid Percentages", "The total percentage must equal 100%.")
            return

        cognitive_total = sum(spinbox.value() for spinbox in self.cognitive_types.values())
        if cognitive_total != 100:
            QMessageBox.warning(self, "Invalid Cognitive Percentages", "The total cognitive percentage must equal 100%.")
            return

        selected_topic = self.topics_list.selectedItems()[0].text() if self.topics_list.selectedItems() else None
        selected_subtopics = []
        selected_topic_for_sub = []
        for (topic, sub), cb in self.subtopic_checkboxes.items():
            if cb.isChecked():
                selected_subtopics.append(sub)
                selected_topic_for_sub.append(topic)

        if not selected_topic or not selected_subtopics:
            QMessageBox.warning(self, "No Topics Selected", "Please select a topic and at least one subtopic.")
            return

        num_questions = int(self.num_questions_spinbox.value())
        difficulty = self.difficulty_combobox.currentText()
        exam_name = self.exam_name_edit.text().strip() or "Untitled Exam"
        exam_type = self.exam_type_combobox.currentText()

        # --- Load question templates ---
        question_templates_beginner = load_question_templates("question_templates_beginner.json")
        question_templates_intermediate = load_question_templates("question_templates_intermediate.json")
        question_templates_advanced = load_question_templates("question_templates_advanced.json")

        if difficulty.lower() == "beginner":
            question_templates = question_templates_beginner
        elif difficulty.lower() == "intermediate":
            question_templates = question_templates_intermediate
        elif difficulty.lower() == "advanced":
            question_templates = question_templates_advanced
        else:
            question_templates = question_templates_beginner

        qtype_to_json_key = {
            "Multiple Choice": "multiple_choice",
            "Problem Solving": "problem_solving",
            "True/False": "true_false"
        }
        qtype_order = [q for q in ["Multiple Choice", "True/False", "Problem Solving"] if self.question_types[q].value() > 0]

        # --- Cognitive type mapping ---
        cog_type_map = {
            "Remembering": ["remember"],
            "Inferential": ["infer", "stand", "analyz"],
            "Applied": ["appl", "evaluat", "creat", "problem"]
        }
        cognitive_percentages = {k: int(v.value()) for k, v in self.cognitive_types.items()}

        # --- Build pools for each cognitive type ---
        cognitive_pools = {ctype: [] for ctype in cognitive_percentages}
        for topic, subtopic in zip(selected_topic_for_sub, selected_subtopics):
            for qtype in qtype_order:
                json_key = qtype_to_json_key[qtype]
                pool = question_templates.get(topic, {}).get(subtopic, {}).get(json_key, [])
                for question in pool:
                    q_type = question.get("type", "").lower()
                    if q_type in ("remembering", "inferential", "applied"):
                        cognitive_pools[q_type].append({
                            "type": question.get("type", qtype),
                            "question": question["question"],
                            "choices": question.get("choices"),
                            "answer": question.get("answer"),
                            "subtopic": subtopic,
                            "topic": topic,
                            "class": json_key
                        })
                        break

        # Shuffle each pool for randomness
        for pool in cognitive_pools.values():
            random.shuffle(pool)

        # --- Calculate exact number of questions per cognitive type ---
        cog_type_counts = {}
        remainder = num_questions
        cog_types = list(cognitive_percentages.keys())
        for i, ctype in enumerate(cog_types):
            if i < len(cog_types) - 1:
                count = round(num_questions * cognitive_percentages[ctype] / 100)
                cog_type_counts[ctype] = count
                remainder -= count
            else:
                cog_type_counts[ctype] = remainder  # Assign remainder to last type

        # --- Select questions for each cognitive type ---
        exam_questions = []
        used_questions = set()
        for ctype, count in cog_type_counts.items():
            pool = cognitive_pools[ctype]
            selected = 0
            for q in pool:
                qid = (q["question"], tuple(q.get("choices", [])) if q.get("choices") else ())
                if qid not in used_questions:
                    exam_questions.append(q)
                    used_questions.add(qid)
                    selected += 1
                    if selected >= count:
                        break
            if selected < count:
                QMessageBox.warning(
                    self,
                    "Not Enough Questions",
                    f"Not enough '{ctype}' questions available. Needed {count}, found {selected}.\n"
                    "Please adjust your settings or add more questions."
                )
                return

        # Shuffle final exam questions
        random.shuffle(exam_questions)

        # Show the generating animation
        dialog = GenerateExamDialog(self)
        dialog.start_animation()
        dialog.exec()

        n_subtopics = len(selected_subtopics)
        time_distribution = random_time_distribution(n_subtopics, 90) if n_subtopics > 0 else []

        # Save to exam bank
        exam_data = {
            "title": exam_name,
            "date": datetime.date.today().strftime("%Y-%m-%d"),
            "type": exam_type,
            "difficulty": difficulty,
            "question_count": len(exam_questions),
            "questions": exam_questions,
            "selected_topics": list(set(selected_topic_for_sub)),
            "selected_subtopics": selected_subtopics,
            "selected_topic_subtopic_pairs": list(zip(selected_topic_for_sub, selected_subtopics)),
            "time_spent_distribution": time_distribution
        }
        self.exam_bank_window.exams.append(exam_data)
        self.exam_bank_window.save_exams()
        self.exam_bank_window.update_exam_slots()
        QMessageBox.information(self, "Exam Generated", "Your exam has been successfully generated!")
        LogsWindow.add_log("Admin", f"Created exam '{exam_name}'")
        self.reset_form()

def random_time_distribution(n, total):
    if n <= 1:
        return [total]
    import random
    cuts = sorted(random.sample(range(1, total), n-1))
    return [cuts[0]] + [cuts[i] - cuts[i-1] for i in range(1, n-1)] + [total - cuts[-1]]

class GenerateExamDialog(QDialog):
    def __init__(self, parent=None, message="Processing..."):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setFixedSize(400, 200)

        # Main layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(10)

        # Custom message label
        self.generating_label = QLabel(message)
        self.generating_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.generating_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #00FF66;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
        """)
        layout.addWidget(self.generating_label)

        # Loading bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                background-color: #eee;
                border: 2px solid #bbb;
                border-radius: 10px;
                height: 20px;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                border-radius: 10px;
            }
        """)
        layout.addWidget(self.progress_bar)

        # Timer for animations
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_progress)

        # Animation state
        self.progress = 0

    def start_animation(self):
        """Start the loading animation."""
        self.progress = 0
        self.progress_bar.setValue(0)
        self.timer.start(50)  # Loading bar animation speed

    def update_progress(self):
        """Update the loading bar progress."""
        if self.progress < 100:
            self.progress += 1
            self.progress_bar.setValue(self.progress)
        else:
            self.timer.stop()
            self.show_completed_message()

    def show_completed_message(self):
        """Show the 'Completed' message."""
        self.accept()  # Close the current dialog
        success_dialog = QDialog(self.parent())
        success_dialog.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        success_dialog.setFixedSize(400, 200)

        # Layout for the success dialog
        layout = QVBoxLayout(success_dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(10)

        # Success message
        success_label = QLabel("Succesfully Done!")
        success_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        success_label.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #00FF66;  # Changed from white to green
            }
        """)
        layout.addWidget(success_label)

        # OK button
        ok_button = QPushButton("OK")
        ok_button.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45A049;
            }
        """)
        ok_button.clicked.connect(success_dialog.accept)
        layout.addWidget(ok_button, alignment=Qt.AlignmentFlag.AlignCenter)

        success_dialog.exec()

def export_exam_to_word(title, exam_type, questions, output_path):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import os

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(14)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # --- Header table: logo left, header text perfectly centered, right cell empty ---
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.3)  # Logo
    table.columns[1].width = Inches(4.8)  # Center cell (centered header)
    table.columns[2].width = Inches(1.3)  # Empty (balances logo)

    # Set row height to at least 1.09"
    tbl = table._tbl
    tr = tbl.tr_lst[0]
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(1.09 * 1440)))  # 1.09 inches in twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Logo cell (top-aligned)
    logo_path = "C:/Users/Administrator/Downloads/PhyGen/logo.png"
    cell_logo = table.cell(0, 0)
    cell_logo.vertical_alignment = 0  # 0 = TOP
    if os.path.exists(logo_path):
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.1))

    # Header text cell (top-aligned, centered horizontally)
    cell_header = table.cell(0, 1)
    cell_header.vertical_alignment = 0  # 0 = TOP
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line, bold in [
        ("Republic of the Philippines", False),
        ('EULOGIO "AMANG" RODRIGUEZ', True),
        ("INSTITUTE OF SCIENCE AND TECHNOLOGY", True),
        ("Nagtahan, Sampaloc Manila", False)
    ]:
        run = p.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.bold = bold

    # Remove table borders for a clean look
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for line in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{line}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    doc.add_paragraph("")

    # --- Exam name and type (centered, below header) ---
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(12)
    t.runs[0].font.name = "Arial"
    t.runs[0].bold = True

    ty = doc.add_paragraph(exam_type)
    ty.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ty.runs[0].font.size = Pt(10)
    ty.runs[0].font.name = "Arial"
    ty.runs[0].bold = True

    # Student info (2x2 table, proper layout)
    info_table = doc.add_table(rows=2, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Inches(4)
    info_table.columns[1].width = Inches(3.4)
    info_table.cell(0, 0).text = "Name:________________________________________"
    info_table.cell(0, 1).text = "Date:________________________________________"
    info_table.cell(1, 0).text = "Course/Year/Section:____________________________"
    info_table.cell(1, 1).text = "Score:_______________________________________"
    for row in info_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

    # Directions
    d = doc.add_paragraph(
        "Directions: Read and answer the questions carefully. Write the letter of the correct answer on the space provided before the number."
    )
    d.alignment = 0
    d.runs[0].font.size = Pt(10)
    d.runs[0].font.name = "Arial"

    questions_per_column_first_page = 6
    questions_per_column_other_pages = 7
    num_columns = 2

    total_questions = len(questions)
    page_number = 1
    if total_questions <= questions_per_column_first_page * num_columns:
        total_pages = 1
    else:
        total_pages = 1 + ((total_questions - questions_per_column_first_page * num_columns + questions_per_column_other_pages * num_columns - 1) // (questions_per_column_other_pages * num_columns))

    q_idx = 0
    while q_idx < total_questions:
        if page_number == 1:
            questions_per_column = questions_per_column_first_page
        else:
            questions_per_column = questions_per_column_other_pages

        if page_number == 1:
            left_start = 0
        else:
            left_start = questions_per_column_first_page * num_columns + (page_number - 2) * questions_per_column_other_pages * num_columns
        right_start = left_start + questions_per_column

        table = doc.add_table(rows=questions_per_column, cols=2)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(3.7)
        table.columns[1].width = Inches(3.7)

        # Fill left column
        for row in range(questions_per_column):
            q_idx_left = left_start + row
            if q_idx_left >= total_questions:
                break
            cell = table.cell(row, 0)
            para = cell.paragraphs[0]
            run = para.add_run(f"___{q_idx_left+1}. {questions[q_idx_left]['question']}")
            run.font.size = Pt(10)
            run.font.name = "Arial"
            if questions[q_idx_left].get("choices"):
                for i, choice in enumerate(questions[q_idx_left]["choices"]):
                    choice_run = para.add_run(f"\n    {chr(65+i)}. {choice}")
                    choice_run.font.size = Pt(10)
                    choice_run.font.name = "Arial"

        # Fill right column
        for row in range(questions_per_column):
            q_idx_right = right_start + row
            if q_idx_right >= total_questions:
                break
            cell = table.cell(row, 1)
            para = cell.paragraphs[0]
            run = para.add_run(f"___{q_idx_right+1}. {questions[q_idx_right]['question']}")
            run.font.size = Pt(10)
            run.font.name = "Arial"
            if questions[q_idx_right].get("choices"):
                for i, choice in enumerate(questions[q_idx_right]["choices"]):
                    choice_run = para.add_run(f"\n    {chr(65+i)}. {choice}")
                    choice_run.font.size = Pt(10)
                    choice_run.font.name = "Arial"

        # Footer (only set once)
        if page_number == 1:
            footer = section.footer
            footer_table = footer.add_table(rows=1, cols=3, width=Inches(7.4))
            footer_table.autofit = False
            footer_table.columns[0].width = Inches(2.7)
            footer_table.columns[1].width = Inches(2.0)
            footer_table.columns[2].width = Inches(2.7)
            left_cell = footer_table.cell(0, 0)
            left_cell.text = ""
            center_cell = footer_table.cell(0, 1)
            center_para = center_cell.paragraphs[0]
            center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            center_para.add_run("Page ")
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            r_element = center_para.add_run()._r
            r_element.append(fldChar1)
            r_element.append(instrText)
            r_element.append(fldChar2)
            center_para.add_run(" out of ")
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = "NUMPAGES"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            r_element = center_para.add_run()._r
            r_element.append(fldChar1)
            r_element.append(instrText)
            r_element.append(fldChar2)
            right_cell = footer_table.cell(0, 2)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_run = right_para.add_run("EARIST-QSF-INST-023")
            right_run.font.size = Pt(9)
            right_run.font.name = "Arial"

        page_number += 1
        if (page_number == 2 and total_questions > questions_per_column_first_page * num_columns) or (q_idx_left + questions_per_column * 2 < total_questions):
            doc.add_page_break()
        if page_number == 2:
            q_idx = questions_per_column_first_page * num_columns
        else:
            q_idx += questions_per_column * num_columns

    # --- Add signature blocks on the last page ---
    doc.add_paragraph("\n\n")
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.2)
    sig_table.columns[1].width = Inches(3.2)

    # Prepared by
    sig_table.cell(0, 0).text = "Prepared by:\n" + "_" * 30
    # Checked by
    sig_table.cell(0, 1).text = "Checked by:\n" + "_" * 30

    sig_table.cell(1, 0).text = ""
    sig_table.cell(1, 1).text = ""

    # Approved by (centered, merged)
    a_cell = sig_table.cell(2, 0)
    b_cell = sig_table.cell(2, 1)
    a_cell.merge(b_cell)
    par_approved = a_cell.paragraphs[0]
    par_approved.add_run("Approved by:\n" + "_" * 30)
    par_approved.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    doc.save(output_path)

def export_answer_key_to_word(title, questions, output_path):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import os

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(14)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    # --- Header table: logo left, header text perfectly centered, right cell empty ---
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.3)  # Logo
    table.columns[1].width = Inches(4.8)  # Center cell (centered header)
    table.columns[2].width = Inches(1.3)  # Empty (balances logo)

    # Set row height to at least 1.09"
    tbl = table._tbl
    tr = tbl.tr_lst[0]
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(1.09 * 1440)))  # 1.09 inches in twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Logo cell (top-aligned)
    logo_path = "logo.png"
    cell_logo = table.cell(0, 0)
    cell_logo.vertical_alignment = 0  # 0 = TOP
    if os.path.exists(logo_path):
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.1))

    # Header text cell (top-aligned, centered horizontally)
    cell_header = table.cell(0, 1)
    cell_header.vertical_alignment = 0  # 0 = TOP
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line, bold in [
        ("Republic of the Philippines", False),
        ('EULOGIO "AMANG" RODRIGUEZ', True),
        ("INSTITUTE OF SCIENCE AND TECHNOLOGY", True),
        ("Nagtahan, Sampaloc Manila", False)
    ]:
        run = p.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.bold = bold

    # Remove table borders for a clean look
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for line in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{line}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    doc.add_paragraph("")  # Spacing

    # --- Section Titles ---
    for text, size, bold in [
        ("COLLEGE ARTS AND SCIENCES", 12, True),
        ("KEY TO CORRECTION", 12, True),
        (title, 12, True)
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial Narrow"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')
        run.bold = bold

    doc.add_paragraph("")  # Spacing

    # --- Answer Key Table (2 columns) ---

    total = len(questions)
    n_rows = (total + 1) // 2
    table = doc.add_table(rows=n_rows, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    for row in range(n_rows):
        # Left column
        idx_left = row
        if idx_left < total:
            q = questions[idx_left]
            answer = q.get("answer", "")
            if q.get("choices"):
                try:
                    answer_idx = next(
                        i for i, choice in enumerate(q["choices"])
                        if choice.strip().lower() == answer.strip().lower()
                    )
                    letter = chr(65 + answer_idx)
                except Exception:
                    letter = "?"

                text = f"{idx_left+1}. {letter}"
            else:
                text = f"{idx_left+1}. {answer}"
            cell = table.cell(row, 0)
            para = cell.paragraphs[0]

            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "Arial Narrow"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')
        # Right column
        idx_right = row + n_rows
        if idx_right < total:
            q = questions[idx_right]
            answer = q.get("answer", "")
            if q.get("choices"):
                try:
                    answer_idx = next(
                        i for i, choice in enumerate(q["choices"])
                        if choice.strip().lower() == answer.strip().lower()
                    )
                    letter = chr(65 + answer_idx)
                except Exception:
                    letter = "?"

                text = f"{idx_right+1}. {letter}"
            else:
                text = f"{idx_right+1}. {answer}"
            cell = table.cell(row, 1)
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "Arial Narrow"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')

    # --- Professors' names at bottom right above footer ---
    doc.add_paragraph("")  # Spacing

    # Add a right-aligned block for signatures
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in [
        "Prepared by:",
        "Mr. REENIER R. LEDESMA",
        "Faculty",
        "",
        "Checked by:",
        "JINAMARLYN B. DOCTOR DPA, RPm",
        "Department Head",
        "",
        "Approved by:",
        "Prof. RODORA T. OLIVEROS M.Sc",
        "Dean"
    ]:
        run = p.add_run(line + "\n")
        run.font.size = Pt(12)
        run.font.name = "Arial Narrow"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')
        run.bold = "by:" in line or "Approved" in line or "Checked" in line

    # --- Footer: code at bottom right ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "EARIST-QSF-INST-025"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.runs[0].font.size = Pt(12)
    footer_para.runs[0].font.name = "Arial Narrow"
    footer_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')

    doc.save(output_path)

def extract_text_for_topic(pdf_path, topic, num_pages=5):
    """
    Extracts text related to a topic from a PDF.

    Args:
        pdf_path (str): Path to the PDF file.
        topic (str): The topic to search for in the PDF.
        num_pages (int): Number of pages to extract after finding the topic.

    Returns:
        list: A list of extracted text blocks/questions.
    """
    doc = fitz.open(pdf_path)
    topic_questions = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if topic.lower() in text.lower():
            # Extract the current page and the next few pages
            for offset in range(num_pages):
                if page_num + offset < len(doc):
                    extracted_page = doc[page_num + offset]
                    extracted_text = extracted_page.get_text()
                    topic_questions.extend(extracted_text.split("\n"))
            break # Stop after finding the topic

    # Filter potential questions (e.g., lines starting with numbers)
    questions = [line.strip() for line in topic_questions if line.strip().startswith(tuple(str(i) + "." for i in range(1,  100)))]
    return questions

class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PhyGen - Physics Exam Generator")
        self.setMinimumSize(QSize(800, 600))

        # Create background label first!
        self.background_label = QLabel(self)
        self.background_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        try:
            self.bg_pixmap = QPixmap("C:/Users/Administrator/Downloads/PhyGen/backgroundimage.png")
            if not self.bg_pixmap.isNull():
                self.background_label.setPixmap(self.bg_pixmap.scaled(
                    self.size(),
                    Qt.AspectRatioMode.IgnoreAspectRatio,
                    Qt.TransformationMode.SmoothTransformation
                ))
            else:
                raise FileNotFoundError("Background image not found.")
        except Exception as e:
            print(f"Error loading background image: {e}")
            palette = self.palette()

            palette.setColor(QPalette.ColorRole.Window, QColor(0, 0, 0))
            self.setPalette(palette)
        self.background_label.setGeometry(0, 0, self.width(), self.height())
        self.background_label.lower()

        # Now create central widget and layout
        self.central_widget = QWidget()
        self.central_widget.setStyleSheet("background: transparent;")
        self.setCentralWidget(self.central_widget)
        self.main_layout_login = QVBoxLayout(self.central_widget)
        self.main_layout_login.setContentsMargins(0, 0, 0, 0)

        # Show login screen
        self.show_login_screen()

    def logout(self):
        """Handle user logout and return to the login screen."""
        self.current_user_type = None  # Clear the current user type
        self.central_widget = QWidget()  # Reinitialize the central widget

        self.central_widget.setStyleSheet("background: transparent;")
        self.setCentralWidget(self.central_widget)
        self.main_layout_login = QVBoxLayout(self.central_widget)  # Reinitialize the layout
        self.main_layout_login.setContentsMargins(0, 0, 0, 0)
        self.show_login_screen()  # Return to the login screen
        LogsWindow.add_log(self.current_user_type, "Logged out")

    def show_login_screen(self):
        # Clear the layout
        for i in reversed(range(self.main_layout_login.count())):
            widget = self.main_layout_login.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container_layout = QHBoxLayout(container)
        container_layout.setContentsMargins(50, 50, 50, 50)
        container_layout.setSpacing(50)

        # Login Form
        self.login_form = QFrame()
        self.login_form.setStyleSheet("""
             QFrame {
                background-color: rgba(0, 0, 0, 0);
                border-radius: 0px;
                padding: 0px;
            }
        """)
        form_layout = QVBoxLayout(self.login_form)
        form_layout.setSpacing(15)
        title_label = QLabel("")
        title_label.setStyleSheet(f"""
        QLabel {{
        font-size: {BASE_FONT_SIZE + 12}px;
        font-weight: bold;
        color: white;
        background: transparent;
    }}
""")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        form_layout.addWidget(title_label)

        self.username_lineedit = QLineEdit()
        self.username_lineedit.setPlaceholderText("Username")
        self.username_lineedit.setStyleSheet(f"""
            QLineEdit {{
                background-color: rgba(255, 255, 255, 200);
                border: 1px solid rgba(255, 255, 255, 100);
                border-radius: 5px;
                padding: 10px;
                font-size: {BASE_FONT_SIZE}px;
                min-width: 250px;
                color: black;
            }}
        """)
        form_layout.addWidget(self.username_lineedit)

        self.password_lineedit = QLineEdit()
        self.password_lineedit.setPlaceholderText("Password")
        self.password_lineedit.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_lineedit.setStyleSheet(self.username_lineedit.styleSheet())
        form_layout.addWidget(self.password_lineedit)

        login_button = QPushButton("Login")

        login_button.setStyleSheet(COMMON_BUTTON_STYLE % (SECONDARY_BUTTON_COLOR, SECONDARY_BUTTON_HOVER_COLOR))
        login_button.clicked.connect(self.attempt_login)
        self.username_lineedit.returnPressed.connect(login_button.click)
        self.password_lineedit.returnPressed.connect(login_button.click)
        form_layout.addWidget(login_button)
        container_layout.addWidget(self.login_form, 0, Qt.AlignmentFlag.AlignCenter)

        self.main_layout_login.addWidget(container)

        # --- Exit button at bottom right ---
        exit_layout = QHBoxLayout()
        exit_layout.addStretch()
        exit_button = QPushButton("Exit")
        exit_button.setStyleSheet(COMMON_BUTTON_STYLE % (WARNING_BUTTON_COLOR, WARNING_BUTTON_HOVER_COLOR))
        exit_button.setFixedWidth(100)
        exit_button.clicked.connect(QApplication.instance().quit)
        exit_layout.addWidget(exit_button)
        self.main_layout_login.addLayout(exit_layout)

    def resizeEvent(self, event):
        if hasattr(self, 'bg_pixmap') and not self.bg_pixmap.isNull():
            self.background_label.setPixmap(self.bg_pixmap.scaled(
                self.size(),
                Qt.AspectRatioMode.IgnoreAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            ))
            self.background_label.setGeometry(0, 0, self.width(), self.height())
        super().resizeEvent(event)

    def logout(self):
        """Handle user logout and return to the login screen."""
        self.current_user_type = None  # Clear the current user type
        self.central_widget = QWidget()  # Reinitialize the central widget
        self.central_widget.setStyleSheet("background: transparent;")
        self.setCentralWidget(self.central_widget)
        self.main_layout_login = QVBoxLayout(self.central_widget)  # Reinitialize the layout
        self.main_layout_login.setContentsMargins(0, 0, 0, 0)
        self.show_login_screen()  # Return to the login screen
        LogsWindow.add_log(self.current_user_type, "Logged out")

    def show_login_screen(self):
        # Clear the layout
        for i in reversed(range(self.main_layout_login.count())):
            widget = self.main_layout_login.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container_layout = QHBoxLayout(container)
        container_layout.setContentsMargins(50, 50, 50, 50)
        container_layout.setSpacing(50)

        # Login Form
        self.login_form = QFrame()
        self.login_form.setStyleSheet("""
             QFrame {
                background-color: rgba(0, 0, 0, 0);
                border-radius: 0px;
                padding: 0px;
            }
       

        """)
        form_layout = QVBoxLayout(self.login_form)
        form_layout.setSpacing(15)
        title_label = QLabel("")
        title_label.setStyleSheet(f"""
        QLabel {{
        font-size: {BASE_FONT_SIZE + 12}px;
        font-weight: bold;
        color: white;
        background: transparent;
    }}
""")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        form_layout.addWidget(title_label)

        self.username_lineedit = QLineEdit()
        self.username_lineedit.setPlaceholderText("Username")
        self.username_lineedit.setStyleSheet(f"""
            QLineEdit {{
                background-color: rgba(255, 255, 255, 200);
                border: 1px solid rgba(255, 255, 255, 100);
                border-radius: 5px;
                padding: 10px;
                font-size: {BASE_FONT_SIZE}px;
                min-width: 250px;
                color: black;
            }}
        """)
        form_layout.addWidget(self.username_lineedit)

        self.password_lineedit = QLineEdit()
        self.password_lineedit.setPlaceholderText("Password")
        self.password_lineedit.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_lineedit.setStyleSheet(self.username_lineedit.styleSheet())
        form_layout.addWidget(self.password_lineedit)

        login_button = QPushButton("Login")
        login_button.setStyleSheet(COMMON_BUTTON_STYLE % (SECONDARY_BUTTON_COLOR, SECONDARY_BUTTON_HOVER_COLOR))
        login_button.clicked.connect(self.attempt_login)
        self.username_lineedit.returnPressed.connect(login_button.click)
        self.password_lineedit.returnPressed.connect(login_button.click)
        form_layout.addWidget(login_button)
        container_layout.addWidget(self.login_form, 0, Qt.AlignmentFlag.AlignCenter)

        self.main_layout_login.addWidget(container)

        # --- Exit button at bottom right ---
        exit_layout = QHBoxLayout()
        exit_layout.addStretch()
        exit_button = QPushButton("Exit")
        exit_button.setStyleSheet(COMMON_BUTTON_STYLE % (WARNING_BUTTON_COLOR, WARNING_BUTTON_HOVER_COLOR))
        exit_button.setFixedWidth(100)
        exit_button.clicked.connect(QApplication.instance().quit)
        exit_layout.addWidget(exit_button)
        self.main_layout_login.addLayout(exit_layout)

    def attempt_login(self):
        username = self.username_lineedit.text()
        password = self.password_lineedit.text()
        hashed_password = hashlib.md5(password.encode()).hexdigest()
        print(f"Entered Username: {username}")
        print(f"Entered Password (hashed): {hashed_password}")



        role = validate_user(username, password)
        # Only accept string or tuple, not bool
        if isinstance(role, tuple):
            self.current_user_type = role[0]
        elif isinstance(role, str):
            self.current_user_type = role
        else:
            self.current_user_type = None

        if self.current_user_type:
            print(f"Login successful! Role: {self.current_user_type}")
            LogsWindow.add_log(username, "Logged in")
            self.show_main_app()
        else:
            print("Login failed: Invalid username or password")
            QMessageBox.critical(
                self,
                "Login Failed",
                "Invalid username or password"
            )

    def show_main_app(self):
        # Clear the login screen layout
        for i in reversed(range(self.main_layout_login.count())):
            widget = self.main_layout_login.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        self.main_app_widget = QWidget()
        self.main_layout_app = QHBoxLayout(self.main_app_widget)
        self.main_layout_app.setContentsMargins(0, 0, 0, 0)
        self.main_layout_app.setSpacing(0)

        # Sidebar
        self.sidebar_frame = QFrame()
        self.sidebar_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(40, 40, 40, 200);
                border-top-right-radius: 20px;
                border-bottom-right-radius: 20px;
                border: none;
                min-width: 200px;
                max-width: 200px;
            }
        """)
        self.sidebar_layout = QVBoxLayout(self.sidebar_frame)
        self.sidebar_layout.setContentsMargins(10, 20, 10, 20)
        self.sidebar_layout.setSpacing(10)

        sidebar_header_btn = QPushButton("PhyGen")
        sidebar_header_btn.setStyleSheet(f"""
    QPushButton {{
        color: white;
        background: transparent;
        border: none;
        padding: 10px;
        text-align: left;
        font-size: 38px;
        font-weight: bold;
        letter-spacing: 2px;
    }}
    QPushButton:hover {{
        color: #00FF66;
        text-decoration: underline;
        background: transparent;
    }}
""")
        sidebar_header_btn.clicked.connect(lambda: self.content_stacked_widget.setCurrentWidget(self.admin_dashboard_widget))
        self.sidebar_layout.addWidget(sidebar_header_btn)

        separator_frame = QFrame()
        separator_frame.setFrameShape(QFrame.Shape.HLine)
        separator_frame.setStyleSheet("color: rgba(255, 255, 255, 100);")
        self.sidebar_layout.addWidget(separator_frame)

        # Navigation buttons
        nav_buttons = [
            ("New Exam", self.show_new_exam_window),
            ("Exam Bank", self.show_exam_bank_window),
            ("Logs", self.show_logs_window),
        ]

        # Add admin-only features
        if self.current_user_type == "admin":
            nav_buttons.append(("Test Bank", self.show_test_bank_window))
            nav_buttons.append(("Generate New User", self.show_user_account_generator))

        for text, callback in nav_buttons:
            nav_button = QPushButton(text)
            nav_button.setStyleSheet(f"""
                QPushButton {{
                    background-color: rgba(60, 60, 60, 150);
                    color: #4CAF50;
                    border: none;
                    border-radius:  10px;
                    padding: 12px;
                    font-size: {BASE_FONT_SIZE}px;
                    text-align: left;
                    padding-left: 20px;
                    font-family: 'Segoe UI', Arial, sans-serif;
                }}
                QPushButton:hover {{
                    background-color: rgba(80, 80, 80, 200);
                }}
            """)
            nav_button.clicked.connect(callback)
            self.sidebar_layout.addWidget(nav_button)

        self.sidebar_layout.addStretch()

        # Dynamically set the "Logged in as" label
        user_info_label = QLabel(f"Logged in as: {self.current_user_type.capitalize()}")
        user_info_label.setStyleSheet(f"""
            QLabel {{
                color: rgba(200, 200, 200, 200);
                font-size: {BASE_FONT_SIZE}px;
                padding: 5px;
                text-align: center;
            }}
        """)
        self.sidebar_layout.addWidget(user_info_label)

        logout_button = QPushButton("Logout")
        logout_button.setStyleSheet(COMMON_BUTTON_STYLE % (WARNING_BUTTON_COLOR, WARNING_BUTTON_HOVER_COLOR))
        logout_button.clicked.connect(self.logout)
        self.sidebar_layout.addWidget(logout_button)

        self.main_layout_app.addWidget(self.sidebar_frame)

        # Content Area
        self.content_stacked_widget = QStackedWidget()
        self.content_stacked_widget.setStyleSheet(f"""
            QStackedWidget {{
                background: 
                    qlineargradient(
                        x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(0,0,0,0.7),
                        stop:0.15 rgba(0,0,0,0.2),
                        stop:0.85 rgba(0,0,0,0.2),
                        stop:1 rgba(0,0,0,0.7)
                    ),
                    url('44.png');
                background-position: center;
                background-repeat: no-repeat;
                background-attachment: fixed;
                background-size: cover;


                border-top-left-radius: 20px; 
                border-bottom-left-radius: 20px;
                color: #00FF66;
                font-size: 14px;
                font-family: 'Segoe UI', Arial, sans-serif;
            }}
        """)

        # Initialize all windows
        self.admin_dashboard_widget = AdminDashboard(self, self.current_user_type)
        self.exam_bank_window_widget = ExamBankWindow()
        self.new_exam_window_widget = NewExamWindow(self.exam_bank_window_widget)
        self.test_bank_window_widget = TestBankWindow()
        self.logs_window_widget = LogsWindow()
        self.user_account_generator_widget = UserAccountGenerator()

        # Add widgets to stacked widget
        self.content_stacked_widget.addWidget(self.admin_dashboard_widget)
        self.content_stacked_widget.addWidget(self.new_exam_window_widget)
        self.content_stacked_widget.addWidget(self.exam_bank_window_widget)
        self.content_stacked_widget.addWidget(self.test_bank_window_widget)
        self.content_stacked_widget.addWidget(self.logs_window_widget)
        self.content_stacked_widget.addWidget(self.user_account_generator_widget)

        # Set default view to dashboard
        self.content_stacked_widget.setCurrentWidget(self.admin_dashboard_widget)

        self.main_layout_app.addWidget(self.content_stacked_widget)
        self.setCentralWidget(self.main_app_widget)

    def show_new_exam_window(self):
        self.content_stacked_widget.setCurrentWidget(self.new_exam_window_widget)

    def show_exam_bank_window(self):
        self.content_stacked_widget.setCurrentWidget(self.exam_bank_window_widget)

    def show_test_bank_window(self):
        self.content_stacked_widget.setCurrentWidget(self.test_bank_window_widget)

    def show_logs_window(self):
        self.content_stacked_widget.setCurrentWidget(self.logs_window_widget)

    def show_user_account_generator(self):
        self.content_stacked_widget.setCurrentWidget(self.user_account_generator_widget)

    def login(self, username, password):
        # Load users from the database
        with open("database.json", "r") as file:
            data = json.load(file)
            users = data.get("users", [])

        # Check credentials
        for user in users:
            if user["username"] == username and user["password"] == hashlib.md5(password.encode()).hexdigest():
                self.current_user = user
                self.show_main_window()
                return

    def show_main_window(self):
        """Show the main window and update the logged-in user role."""
        self.main_window = QMainWindow(self.current_user_type)  # Pass the role
        self.main_window.show()
        self.close()

class LogsWindow(QMainWindow):
    LOGS_FILE = "logs.json"

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Application Logs")
        self.setMinimumSize(900, 700)

        # Central widget and layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Title
        title_label = QLabel("Application Logs")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 32px;
                font-weight: bold;
                color: #009944;



                text-align: center;
                margin-bottom: 20px;
            }
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        # --- Sort By Row ---
        sort_row = QHBoxLayout()
        sort_label = QLabel("Sort by:")
        sort_label.setStyleSheet("font-size: 22px; color: #009944; font-weight: bold;")
        sort_row.addWidget(sort_label)

        self.sort_combobox = QComboBox()
        self.sort_combobox.addItems(["Latest", "Oldest"])
        self.sort_combobox.setStyleSheet("""
            QComboBox {
                background-color: #3c3c3c;
                color: #009944;
                border: 2px solid #009944;
                border-radius: 8px;
                padding: 8px;
                font-size: 22px;
                min-width: 180px;
            }
        """)
        self.sort_combobox.currentIndexChanged.connect(self.load_logs)
        sort_row.addWidget(self.sort_combobox)
        sort_row.addStretch()
        layout.addLayout(sort_row)

        # Logs table
        self.logs_table = QTableWidget()
        self.logs_table.setColumnCount(4)
        self.logs_table.setHorizontalHeaderLabels(["User", "Activity", "Time", "Date"])
        self.logs_table.setStyleSheet("""
            QTableWidget {
                background-color: #fff;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                font-size: 14px;
                color: #222;
            }
            QHeaderView::section {
                background-color: #F5F5F5;
                color: #222;
                font-weight: normal;
                font-size: 14px;
                border: none;
                padding: 5px;
            }
            QTableWidget::item {
                border: none;
                padding: 5px;
            }
        """)
        self.logs_table.horizontalHeader().setStretchLastSection(True)
        self.logs_table.horizontalHeader().setDefaultAlignment(Qt.AlignmentFlag.AlignCenter)

        self.logs_table.verticalHeader().setVisible(False)
        self.logs_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.logs_table.verticalHeader().setDefaultSectionSize(40)  # You can adjust 40 to your preferred height
        layout.addWidget(self.logs_table)

        # Load logs
        self.load_logs()

    def load_logs(self):
        """Load logs from the logs file and display them in the table."""
       
        if not os.path.exists(self.LOGS_FILE):
            self.logs_table.setRowCount(0)
            return







        with open(self.LOGS_FILE, "r", encoding="utf-8") as file:
            logs = json.load(file)

        # Sort logs
        sort_mode = self.sort_combobox.currentText() if hasattr(self, "sort_combobox") else "Latest"
        if sort_mode == "Latest":
            logs = sorted(logs, key=lambda x: (x.get("date", ""), x.get("time", "")), reverse=True)
        else:
            logs = sorted(logs, key=lambda x: (x.get("date", ""), x.get("time", "")))

        self.logs_table.setRowCount(len(logs))
        for row, log in enumerate(logs):
            self.logs_table.setItem(row, 0, QTableWidgetItem(log.get("user", "")))
            self.logs_table.setItem(row, 1, QTableWidgetItem(log.get("did", "")))
            self.logs_table.setItem(row, 2, QTableWidgetItem(log.get("time", "")))
            self.logs_table.setItem(row, 3, QTableWidgetItem(log.get("date", "")))

    @staticmethod
    def add_log(user, action):
        """Add a new log entry to the logs file."""
        log_entry = {
            "user": user,
            "did": action,
            "time": datetime.datetime.now().strftime("%I:%M %p"),
            "date": datetime.datetime.now().strftime("%m/%d/%Y")
        }

        # Load existing logs
        logs = []
        if os.path.exists(LogsWindow.LOGS_FILE):
            with open(LogsWindow.LOGS_FILE, "r", encoding="utf-8") as file:
                logs = json.load(file)

        # Append the new log entry
        logs.append(log_entry)

        # Save the updated logs back to the file
        with open(LogsWindow.LOGS_FILE, "w", encoding="utf-8") as file:
            json.dump(logs, file, indent=4)
        print(f"Adding log: {log_entry}")

class UserAccountGenerator(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Generate New User")
        self.setMinimumSize(500, 400)

        # Central widget and layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        layout = QVBoxLayout(self.central_widget)

        # Title
        title_label = QLabel("Generate New User")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 28px;
                font-weight: bold;
                color: #222;
                text-align: center;
                margin-bottom: 20px;
                letter-spacing: 2px;
                background: transparent;
            }
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        # Form container
        form_container = QFrame()
        form_container.setStyleSheet("""
            QFrame {
                background: transparent;
                border: 1.5px solid #222;
                border-radius: 12px;
                padding: 28px;
            }
        """)
        form_layout = QVBoxLayout(form_container)
        form_layout.setSpacing(18)

        # Modern black/white fields
        field_style = """
            QLineEdit {
                background: #000;
                color: #fff;
                border: 1px solid #222;
                border-radius: 6px;
                padding: 12px;
                font-size: 16px;
            }
            QLineEdit:focus {
                border: 1.5px solid #222;
                background: #181818;
            }
        """

        # Full Name
        self.full_name_edit = QLineEdit()
        self.full_name_edit.setPlaceholderText("Full Name")
        self.full_name_edit.setStyleSheet(field_style)
        form_layout.addWidget(self.full_name_edit)

        # Username
        self.username_edit = QLineEdit()
        self.username_edit.setPlaceholderText("Username")
        self.username_edit.setStyleSheet(field_style)
        form_layout.addWidget(self.username_edit)

        # Password
        self.password_edit = QLineEdit()
        self.password_edit.setPlaceholderText("Password")
        self.password_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_edit.setStyleSheet(field_style)
        form_layout.addWidget(self.password_edit)

        # Confirm Password
        self.confirm_password_edit = QLineEdit()
        self.confirm_password_edit.setPlaceholderText("Confirm Password")
        self.confirm_password_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.confirm_password_edit.setStyleSheet(field_style)
        form_layout.addWidget(self.confirm_password_edit)

        layout.addWidget(form_container)

        # Generate Button
        generate_button = QPushButton("Generate Account")
        generate_button.setStyleSheet("""
            QPushButton {
                background-color: #fff;
                color: #222;
                border: 1.5px solid #222;
                border-radius: 6px;
                padding: 12px;
                font-size: 16px;
                font-weight: bold;
                letter-spacing: 1px;
            }
            QPushButton:hover {
                background-color: #222;
                color: #fff;
                border: 1.5px solid #222;
            }
        """)
        generate_button.clicked.connect(self.generate_user_account)
        layout.addWidget(generate_button, 0, Qt.AlignmentFlag.AlignCenter)

    def generate_user_account(self):
        username = self.username_edit.text().strip()
        password = self.password_edit.text()
        confirm_password = self.confirm_password_edit.text()
        full_name = self.full_name_edit.text().strip()

        if not username or not password or password != confirm_password:
            QMessageBox.warning(self, "Validation Error", "Invalid input")
            return

        # Show the "Creating Account..." animation
        dialog = GenerateExamDialog(self, message="Creating Account...")
        dialog.start_animation()
        dialog.exec()

        # Add the user to the database
        with open("database.json", "r") as file:
            data = json.load(file)

        hashed_password = hashlib.md5(password.encode()).hexdigest()
        data["users"].append({
            "username": username,
            "password": hashed_password,
            "role": "non-admin"
        })

        with open("database.json", "w") as file:
            json.dump(data, file, indent=4)

        QMessageBox.information(self, "Success", f"User '{username}' created successfully.")

class TestBankWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Test Bank")
        self.setMinimumSize(900, 700)
        self.tests = []
        self.questions_data = {}
        self.current_difficulty = "Beginner"
        self.file_map = {
            "Beginner": "question_templates_beginner.json",
            "Intermediate": "question_templates_intermediate.json",
            "Advanced": "question_templates_advanced.json"
        }

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Controls
        top_row_layout = QHBoxLayout()
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Search...")
        self.search_box.setFixedWidth(180)
        self.search_box.textChanged.connect(self.filter_questions)
        top_row_layout.addWidget(self.search_box)

        self.difficulty_dropdown = QComboBox()
        self.difficulty_dropdown.addItems(["Beginner", "Intermediate", "Advanced"])
        self.difficulty_dropdown.currentTextChanged.connect(self.on_difficulty_changed)
        top_row_layout.addWidget(self.difficulty_dropdown)

        add_button = QPushButton("Add")
        add_button.setFixedWidth(80)
        add_button.clicked.connect(self.add_question)
        top_row_layout.addWidget(add_button)

        delete_button = QPushButton("Delete")
        delete_button.setFixedWidth(80)
        delete_button.clicked.connect(self.delete_selected_questions)
        top_row_layout.addWidget(delete_button)

        import_button = QPushButton("Import")
        import_button.setFixedWidth(80)
        import_button.clicked.connect(self.import_questions)
        top_row_layout.addWidget(import_button)

        top_row_layout.addStretch()
        layout.addLayout(top_row_layout)

        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(7)
        self.table.setHorizontalHeaderLabels(["Topic", "Subtopic", "Type", "Class", "Question", "Choices", "Answer"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.verticalHeader().setVisible(False)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        layout.addWidget(self.table)

        self.ensure_files_exist()
        # Force loading of the current difficulty on startup
        self.difficulty_dropdown.setCurrentText("Beginner")
        self.load_questions()

    def ensure_files_exist(self):
        """Ensure all question template files exist and are valid JSON."""
        for path in self.file_map.values():
            if not os.path.exists(path):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump({}, f)
            else:
                # If file is empty or invalid, reset to {}
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        if not isinstance(data, dict):
                            raise ValueError
                except Exception:
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump({}, f)

    def on_difficulty_changed(self, difficulty):
        self.current_difficulty = difficulty
        self.load_questions()

    def load_questions(self):
        """Load questions from the current difficulty's JSON file into the table and self.tests."""
        self.table.setRowCount(0)
        difficulty = self.current_difficulty
        file_path = self.file_map[difficulty]
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # ...inside TestBankWindow.load_questions()...
        for topic, subtopics in data.items():
            for subtopic, classes in subtopics.items():
                for qclass, qlist in classes.items():
                    for q in qlist:
                        # Use the actual class from the entry if present, else fallback to qclass
                        real_class = qclass
                        self.tests.append({
                            "topic": topic,
                            "subtopic": subtopic,
                            "type": q.get("type", ""),
                            "class": real_class,
                            "question": q.get("question", ""),
                            "choices": q.get("choices", []),
                            "answer": q.get("answer", "")
                        })
                        # Add row to table
                        row = self.table.rowCount()
                        self.table.insertRow(row)
                        self.table.setItem(row, 0, QTableWidgetItem(topic))
                        self.table.setItem(row, 1, QTableWidgetItem(subtopic))
                        self.table.setItem(row, 2, QTableWidgetItem(q.get("type", "")))
                        self.table.setItem(row, 3, QTableWidgetItem(real_class.replace("_", " ").title()))
                        self.table.setItem(row, 4, QTableWidgetItem(q.get("question", "")))  # Question column
                        self.table.setItem(row, 5, QTableWidgetItem(", ".join(q.get("choices", [])) if q.get("choices") else ""))
                        self.table.setItem(row, 6, QTableWidgetItem(q.get("answer", "")))    # Answer column

    def save_questions(self):
        """Save the current questions to the JSON file, grouped by type."""
        file_path = self.file_map[self.current_difficulty]
        output = {}
        for q in self.tests:
            topic = q["topic"]
            subtopic = q["subtopic"]
            qtype = q["type"]
            output.setdefault(topic, {})
            output[topic].setdefault(subtopic, {})
            qclass = q.get("class", "multiple_choice")
            output[topic][subtopic].setdefault(qclass, [])
            entry = {
                "question": q["question"],
                "answer": q["answer"],
                "type": q.get("type", ""),
            }
            if q.get("choices"):
                entry["choices"] = q["choices"]
            output[topic][subtopic][qclass].append(entry)
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=4)
        self.load_questions()  # Always reload after saving
        # Update dashboard stats if open
        main_window = self.parent()
        if main_window and hasattr(main_window, "admin_dashboard_widget"):
            main_window.admin_dashboard_widget.update_stats()

        if self.parent() and hasattr(self.parent(), "admin_dashboard_widget"):
            self.parent().admin_dashboard_widget.update_stats()

    def populate_table(self, questions):
        """Populate the table with the given list of questions."""
        self.table.setRowCount(0)
        for q in questions:
            row = self.table.rowCount()
            self.table.insertRow(row)
            self.table.setItem(row, 0, QTableWidgetItem(q["topic"]))
            self.table.setItem(row, 1, QTableWidgetItem(q["subtopic"]))
            self.table.setItem(row, 2, QTableWidgetItem(q["type"]))
            self.table.setItem(row, 3, QTableWidgetItem(q["class"].replace(" ", "_").lower()))  # Class column
            self.table.setItem(row, 4, QTableWidgetItem(q["question"]))  # Question column
            self.table.setItem(row, 5, QTableWidgetItem(q["answer"]))    # Answer column

    def filter_questions(self):
        """Filter questions so only those with topics starting with the search text are shown."""
        text = self.search_box.text().strip().lower()
        if not text:
            self.populate_table(self.tests)
        else:
            filtered = [
                q for q in self.tests
                if q["topic"].lower().startswith(text)
            ]
            self.populate_table(filtered)

    def add_question(self):
        """Add a new question to the test bank, with type selection."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Question")
        layout = QVBoxLayout(dialog)

        topic_input = QLineEdit()
        subtopic_input = QLineEdit()
        type_combo = QComboBox()
        type_combo.addItems(["remembering", "inferential", "applied"])
        class_combo = QComboBox()
        class_combo.addItems(["Multiple Choice", "True/False", "Problem Solving"])
        question_input = QTextEdit()
        answer_input = QLineEdit()
        choices_label = QLabel("Choices:")
        choices_input = QTextEdit()
        choices_input.setPlaceholderText("Enter choices (one per line)")
        layout.addWidget(QLabel("Topic:"))
        layout.addWidget(topic_input)
        layout.addWidget(QLabel("Subtopic:"))
        layout.addWidget(subtopic_input)
        layout.addWidget(QLabel("Type:"))
        layout.addWidget(type_combo)
        layout.addWidget(QLabel("Class:"))
        layout.addWidget(class_combo)
        layout.addWidget(QLabel("Question:"))
        layout.addWidget(question_input)
        layout.addWidget(choices_label)
        layout.addWidget(choices_input)
        layout.addWidget(QLabel("Answer:"))
        layout.addWidget(answer_input)

        def update_choices_visibility():
            is_mc = class_combo.currentText() == "Multiple Choice"
            choices_label.setVisible(is_mc)
            choices_input.setVisible(is_mc)
        class_combo.currentIndexChanged.connect(update_choices_visibility)
        update_choices_visibility()

        btns = QHBoxLayout()
        add_btn = QPushButton("Add")
        cancel_btn = QPushButton("Cancel")
        btns.addWidget(add_btn)
        btns.addWidget(cancel_btn)
        layout.addLayout(btns)

        def on_add():
            topic = topic_input.text().strip()
            subtopic = subtopic_input.text().strip()
            qtype = type_combo.currentText()
            qclass = class_combo.currentText()
            question = question_input.toPlainText().strip()
            answer = answer_input.text().strip()
            choices = [c.strip() for c in choices_input.toPlainText().splitlines() if c.strip()] if qclass == "Multiple Choice" else None

            # Map display class to JSON key
            class_map = {
                "Multiple Choice": "multiple_choice",
                "True/False": "true_false",
                "Problem Solving": "problem_solving"
            }
            json_class = class_map.get(qclass, "multiple_choice")

            entry = {
                "question": question,
                "answer": answer,
                "type": qtype
            }
            if choices:
                entry["choices"] = choices

            # --- Load the correct JSON file for current difficulty ---
            difficulty = self.current_difficulty
            file_path = self.file_map[difficulty]
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            # --- Ensure topic/subtopic and arrays exist ---
            if topic not in data:
                data[topic] = {}
            if subtopic:
                if subtopic not in data[topic]:
                    data[topic][subtopic] = {k: [] for k in ["multiple_choice", "true_false", "problem_solving"]}
                target = data[topic][subtopic]
            else:
                if not isinstance(data[topic], dict) or not all(k in data[topic] for k in ["multiple_choice", "true_false", "problem_solving"]):
                    data[topic] = {k: [] for k in ["multiple_choice", "true_false", "problem_solving"]}
                target = data[topic]

            # --- Add to the correct array only ---
            target[json_class].append(entry)

            # --- Save back to JSON file ---
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)

            dialog.accept()
            self.load_questions()  # <-- This will refresh the table!

        add_btn.clicked.connect(on_add)
        cancel_btn.clicked.connect(dialog.reject)
        dialog.exec()

    def delete_selected_questions(self):
        if get_admin_passcode() and not require_admin_passcode(self):
            return
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Delete", "Please select a question to delete.")
            return

        # Get the question data from the table row (not from self.tests index)
        topic = self.table.item(selected, 0).text()
        subtopic = self.table.item(selected, 1).text()
        qtype = self.table.item(selected, 2).text()
        qclass = self.table.item(selected, 3).text().replace(" ", "_").lower()
        question = self.table.item(selected, 4).text()
        choices = self.table.item(selected, 5).text()
        answer = self.table.item(selected, 6).text()

        # Show a more informative preview
        preview = (
            f"Topic: {topic}\n"
            f"Subtopic: {subtopic}\n"
            f"Type: {qtype}\n"
            f"Class: {qclass}\n"
            f"Question: {question}\n"
            f"Choices: {choices}\n"
            f"Answer: {answer}"
        )
        confirm = QMessageBox.question(
            self, "Confirm Delete", f"Delete this question?\n\n{preview}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if confirm != QMessageBox.StandardButton.Yes:
            return

        # Remove from JSON file
        file_path = self.file_map[self.current_difficulty]
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Find and remove the question from the JSON structure
        found = False
        if topic in data and subtopic in data[topic] and qclass in data[topic][subtopic]:
            new_list = []
            for q in data[topic][subtopic][qclass]:
                q_choices = ", ".join(q.get("choices", [])) if q.get("choices") else ""
                if (
                    q.get("question", "") == question and
                    q.get("answer", "") == answer and
                    q_choices == choices
                ):
                    found = True
                    continue  # skip this question (delete)
                new_list.append(q)
            data[topic][subtopic][qclass] = new_list

        if found:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            self.load_questions()
            LogsWindow.add_log("admin", f"Deleted question: {question}")
        else:
            QMessageBox.warning(self, "Delete", "Could not find the question in the database.")

    def import_questions(self):
        """Import questions from a formatted .txt or .docx file."""
        # 1. Ask for difficulty
        difficulties = ["Beginner", "Intermediate", "Advanced"]
        diff, ok = QInputDialog.getItem(self, "Select Difficulty", "Difficulty:", difficulties, 0, False)
        if not ok or not diff:
            return
        self.current_difficulty = diff
        self.difficulty_dropdown.setCurrentText(diff)

        # 2. File dialog
        file_path, _ = QFileDialog.getOpenFileName(self, "Import Questions", "", "Text Files (*.txt);;Word Files (*.docx)")
        if not file_path:
            return

        # 3. Parse file
        def parse_blocks(lines):
            blocks = []
            block = {}
            for line in lines:
                line = line.strip()
                if not line:
                    if block:
                        blocks.append(block)
                        block = {}
                    continue
                if ":" in line:
                    key, val = line.split(":", 1)
                    key = key.strip().lower()
                    val = val.strip()
                    if key == "choices":
                        block[key] = [c.strip() for c in val.split(",") if c.strip()]
                    else:
                        block[key] = val
            if block:
                blocks.append(block)
            return blocks

        import os
        from docx import Document

        try:
            if file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    lines = f.readlines()
            elif file_path.endswith(".docx"):
                doc = Document(file_path)
                lines = []
                for para in doc.paragraphs:
                    lines.append(para.text)
            else:
                QMessageBox.warning(self, "Import Error", "Unsupported file type.")
                return

            questions = parse_blocks(lines)
            if not questions:
                QMessageBox.warning(self, "Import Error", "No valid questions found in file.")
                return

            # 4. Save to JSON
            file_path_json = self.file_map[self.current_difficulty]
            with open(file_path_json, "r", encoding="utf-8") as f:
                data = json.load(f)

            class_map = {
                "multiple choice": "multiple_choice",
                "true/false": "true_false",
                "problem solving": "problem_solving"
            }

            for q in questions:
                topic = q.get("topic", "Unknown")
                subtopic = q.get("subtopic", "Unknown")
                qclass = class_map.get(q.get("class", "").strip().lower(), "multiple_choice")
                entry = {
                    "question": q.get("question", ""),
                    "answer": q.get("answer", ""),
                    "type": q.get("type", ""),
                }
                if "choices" in q:
                    entry["choices"] = q["choices"]

                if topic not in data:
                    data[topic] = {}
                if subtopic not in data[topic]:
                    data[topic][subtopic] = {k: [] for k in class_map.values()}
                data[topic][subtopic][qclass].append(entry)

            with open(file_path_json, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)

            self.load_questions()
            QMessageBox.information(self, "Import Successful", f"{len(questions)} questions imported successfully.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to import questions: {e}")

    def import_exam(self):
        # ... your code to import exam and answer key ...
        # Let's say you parse/import and get:
        imported_subtopics = [...]  # <-- fill this with the real subtopics from the import

        output_path = QFileDialog.getSaveFileName(self, "Save TOS", "TOS.xlsx", "Excel Files (*.xlsx)")[0]
        if not output_path:
            return

        selected_topics = [...]  # fill with actual topics from import
        selected_subtopics = imported_subtopics  # Use all imported subtopics
        num_questions = ...  # actual number from import
        exam_type = ...      # actual value from import
        difficulty = ...     # actual value from import
        question_distribution = ...  # actual distribution from import

        export_tos_to_excel(
            output_path,
            selected_topics,
            selected_subtopics,
            num_questions,
            exam_type,
            difficulty,
            question_distribution,
            faculty="Mr. REENIER R. LEDESMA",
            dept_head="Prof. LESTER D. BERNARDINO M.Sc",
            dean="Prof. RODORA T. OLIVEROS M.Sc",
            questions=self.tests  # Pass the flat list of questions from the current test bank
        )
        QMessageBox.information(self, "TOS Exported", f"TOS file saved as {output_path}")

class SubtopicPreviewDialog(QDialog):
    def __init__(self, checked_subtopics, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Preview Selected Subtopics")
        self.setMinimumSize(400, 300)
        layout = QVBoxLayout(self)
        label = QLabel("You have selected the following subtopics:")
        layout.addWidget(label)
        subtopics_list = QListWidget()
        for (topic, subtopic) in checked_subtopics:
            subtopics_list.addItem(f"{topic} - {subtopic}")
        layout.addWidget(subtopics_list)
        button_layout = QHBoxLayout()
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)
        self.setStyleSheet("QDialog { background-color: #2c2c2c; color: #00FF66; }")

class UserManagementDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Manage Users")
        self.setModal(True)
        self.setMinimumSize(500, 400)
        self.setStyleSheet("""
            QDialog {
                background: rgba(255,255,255,0.95);
            }
        """)
        layout = QVBoxLayout(self)

        # X Button
        close_btn = QPushButton("")
        close_btn.setFixedSize(32, 32)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #222;
                font-size: 20px;
                border: none;
            }
            QPushButton:hover {
                color: #c00;
            }
        """)
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn, alignment=Qt.AlignmentFlag.AlignRight)

        # User list
        self.user_list = QListWidget()
        self.user_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #222;
                background: transparent;
                color: #222;
                font-size: 15px;
            }
            QListWidget::item:selected {
                background: #008000;
                color: #fff;
            }
        """)
        layout.addWidget(self.user_list)

        self.load_users()

        # Delete and View buttons
        btn_layout = QHBoxLayout()
        self.delete_btn = QPushButton("Delete User")
        self.delete_btn.setStyleSheet("background: #fff; color: #c00; border: 1.5px solid #c00; border-radius: 5px;")
        self.delete_btn.clicked.connect(self.delete_user)
        btn_layout.addWidget(self.delete_btn)

        self.view_btn = QPushButton("View User")
        self.view_btn.setStyleSheet("background: #fff; color: #008000; border: 1.5px solid #008000; border-radius: 5px;")
        self.view_btn.clicked.connect(self.view_user)
        btn_layout.addWidget(self.view_btn)

        layout.addLayout(btn_layout)

    def load_users(self):
        self.user_list.clear()
        try:
            with open("database.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            for user in data.get("users", []):
                self.user_list.addItem(user.get("username", ""))
        except Exception:
            pass

    def delete_user(self):
        selected = self.user_list.currentItem()
        if not selected:
            QMessageBox.warning(self, "No Selection", "Select a user to delete.")
            return
        username = selected.text()
        confirm = QMessageBox.question(self, "Confirm Delete", f"Delete user '{username}'?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if confirm == QMessageBox.StandardButton.Yes:
            with open("database.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            data["users"] = [u for u in data.get("users", []) if u.get("username") != username]
            with open("database.json", "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
            self.load_users()

    def view_user(self):
        selected = self.user_list.currentItem()
        if not selected:
            QMessageBox.information(self, "No Selection", "Select a user to view.")
            return
        username = selected.text()
        with open("database.json", "r", encoding="utf-8") as f:
            data = json.load(f)
        user = next((u for u in data.get("users", []) if u.get("username") == username), None)
        if user:
            QMessageBox.information(self, "User Info", f"Username: {user.get('username')}\nRole: {user.get('role', 'N/A')}")

# --- Passcode Utility Functions ---

def get_admin_passcode():
    """Return the admin passcode if set, else None."""
    try:
        with open("database.json", "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("admin_passcode", None)
    except Exception:
        return None

def set_admin_passcode(new_passcode):
    """Set or remove the admin passcode in database.json."""
    try:
        with open("database.json", "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = {}
    if new_passcode:
        data["admin_passcode"] = new_passcode
    else:
        data.pop("admin_passcode", None)
    with open("database.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

def require_admin_passcode(parent=None):
    """Prompt for passcode if set. Returns True if allowed, False otherwise."""
    passcode = get_admin_passcode()
    if not passcode:
        return True  # No passcode set, allow
    entered, ok = QInputDialog.getText(parent, "Admin Passcode Required", "Enter admin passcode:", QLineEdit.EchoMode.Password)
    if ok and entered == passcode:
        return True
    if ok:
        QMessageBox.critical(parent, "Access Denied", "Incorrect passcode.")
    return False

def export_tos_to_excel(
    output_path,
    selected_topics,
    selected_subtopics,
    num_questions,
    exam_type,  # unused
    difficulty,  # unused
    question_distribution,
    faculty="Mr. REENIER R. LEDESMA",
    dept_head="Prof. LESTER D. BERNARDINO M.Sc",
    dean="Prof. RODORA T. OLIVEROS M.Sc",
    questions=None  # <-- Pass the exam's questions list here!
):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    import os

    wb = Workbook()
    ws = wb.active
    ws.title = "TOS"

    col_widths = [5, 32, 15, 20, 10, 10, 10, 10, 10, 10, 10, 10, 10, 40]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    thin = Side(border_style="thin", color="000000")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)

    # --- HEADER LINES ---
    ws.merge_cells('G2:M2')
    ws['G2'] = "Republic of the Philippines"
    ws['G2'].font = Font(size=12, italic=True)
    ws['G2'].alignment = Alignment(horizontal="center")

    ws.merge_cells('G3:M3')
    ws['G3'] = 'EULOGIO "AMANG" RODRIGUEZ'
    ws['G3'].font = Font(size=13, bold=True)
    ws['G3'].alignment = Alignment(horizontal="center")

    ws.merge_cells('G4:M4')
    ws['G4'] = "INSTITUTE OF SCIENCE AND TECHNOLOGY"
    ws['G4'].font = Font(size=12, bold=True)
    ws['G4'].alignment = Alignment(horizontal="center")

    ws.merge_cells('G5:M5')
    ws['G5'] = "Nagtahan, Sampaloc Manila"
    ws['G5'].font = Font(size=11, italic=True)
    ws['G5'].alignment = Alignment(horizontal="center")

    # --- Place logo.png in D3 ---
    logo_path = "logo.png"
    if os.path.exists(logo_path):
        img = XLImage(logo_path)
        img.height = 80
        img.width = 80
        ws.add_image(img, "F2")

    # --- TOS TITLE ---
    ws.merge_cells('A1:N1')
    ws['A1'] = "TABLE OF SPECIFICATIONS"
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal="center")

    # --- SUBJECT PROFESSOR ---
    ws['A4'] = "Subject Professor: Reenier Ledesma"

    # --- Academic Year and Semester moved to C7 and D7 ---
    ws['B7'] = "Academic Year:"
    ws['D7'] = "Semester:"

    ws['A3'] = "Subject:"

    ws.merge_cells('A6:N6')
    ws['F26'] = "Legend: NOI = Number of Items, POI = Placement of Items"
    ws['A6'].font = Font(italic=True, size=10)
    ws['A6'].alignment = Alignment(horizontal="center")

    # Table header (row 8-11)
    ws.merge_cells('A8:A11')
    ws['A8'] = "No."
    ws.merge_cells('B8:B11')
    ws['B8'] = "Subtopic"
    ws.merge_cells('C8:C11')
    ws['C8'] = "Competencies"
    ws.merge_cells('D8:D11')
    ws['D8'] = "Time Spent / Frequency (min)"
    ws.merge_cells('E8:E11')
    ws['E8'] = "Weight %"

    ws.merge_cells('F8:G8')
    ws.merge_cells('H8:I8')
    ws.merge_cells('J8:K8')
    ws['F8'] = "Remembering"
    ws['H8'] = "Inferential"
    ws['J8'] = "Applied"
    ws['F8'].alignment = ws['H8'].alignment = ws['J8'].alignment = Alignment(horizontal="center", vertical="center")
    ws['F8'].font = ws['H8'].font = ws['J8'].font = Font(bold=True)

    ws['F9'] = "NOI"
    ws['G9'] = "POI"
    ws['H9'] = "NOI"
    ws['I9'] = "POI"
    ws['J9'] = "NOI"
    ws['K9'] = "POI"

    for start_col in [6, 8, 10]:
        ws.merge_cells(f"{get_column_letter(start_col)}10:{get_column_letter(start_col+1)}10")
        ws[f"{get_column_letter(start_col)}10"] = ""
        ws[f"{get_column_letter(start_col)}10"].border = border
        ws[f"{get_column_letter(start_col)}10"].alignment = Alignment(horizontal="center", vertical="center")

    ws.merge_cells('L8:M8')
    ws['L8'] = "Total number of test item"
    ws['L8'].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws['L8'].font = Font(bold=True)
    ws.merge_cells('L9:M9')
    ws['L9'] = ""
    ws['L9'].border = border
    ws['L9'].alignment = Alignment(horizontal="center", vertical="center")
    ws['L10'] = "Actual"
    ws['M10'] = "Adjusted"
    ws['L10'].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws['M10'].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws['L10'].font = Font(bold=True)
    ws['M10'].font = Font(bold=True)
    ws.merge_cells('L11:M11')
    ws['L11'] = ""
    ws['L11'].border = border
    ws['L11'].alignment = Alignment(horizontal="center", vertical="center")

    for row in ws.iter_rows(min_row=8, max_row=11, min_col=1, max_col=13):
        for cell in row:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

    # --- Accurate subtopic/type scan using (topic, subtopic) ---
    # Always list all selected subtopics (with topic), in the order selected
    subtopic_pairs = []
    for topic, sub in zip(selected_topics, selected_subtopics):
        subtopic_pairs.append((topic, sub))

    # Build mapping: (topic, subtopic) -> {type -> [question_indices]}
    subtopic_map = {pair: {"remembering": [], "inferential": [], "applied": []} for pair in subtopic_pairs}
    subtopic_total = {pair: 0 for pair in subtopic_pairs}
    total_questions = 0

    # --- Cognitive type mapping (define here for use below) ---
    cog_type_map = {
        "remembering": ["remember"],
        "inferential": ["infer", "stand", "analyz"],
        "applied": ["appl", "evaluat", "creat", "problem"]
    }

    if questions:
        for idx, q in enumerate(questions, 1):
            pair = (q.get("topic", "").strip(), q.get("subtopic", "").strip())
            if pair in subtopic_map:
                q_type = q.get("type", "").strip().lower()
                if q_type in ("remembering", "inferential", "applied"):
                    subtopic_map[pair][q_type].append(idx)
                subtopic_total[pair] += 1
                total_questions += 1
    else:
        total_questions = num_questions
        for pair in subtopic_pairs:
            subtopic_total[pair] = num_questions // len(subtopic_pairs)

    # --- Fair time distribution: at least 15 min per subtopic, all unique, sum to 90 ---
    n_sub = len(subtopic_pairs)
    if n_sub > 0:
        min_time = 15
        total_time = 90
        if n_sub > 5:
            base_time = [min_time] * n_sub
            for i in range(total_time - min_time * n_sub):
                base_time[i % n_sub] += 1
        else:
            unique_times = list(range(min_time, min_time + n_sub))
            current_sum = sum(unique_times)
            diff = total_time - current_sum
            i = 0
            while diff > 0:
                if unique_times[i] + 1 not in unique_times:
                    unique_times[i] += 1
                    diff -= 1
                i = (i + 1) % n_sub
            base_time = unique_times
        import random
        random.shuffle(base_time)
        time_spent = {pair: base_time[i] for i, pair in enumerate(subtopic_pairs)}
    else:
        time_spent = {}

    # --- Weight % is proportional to question count ---
    weight_pct = {}
    for pair in subtopic_pairs:
        w = round((subtopic_total[pair] / total_questions) * 100, 2) if total_questions else 0.0
        weight_pct[pair] = w

    # Adjust rounding errors to ensure total weight is exactly 100%
    total_weight = sum(weight_pct.values())
    if subtopic_pairs and round(total_weight, 2) != 100.0:
        diff = round(100.0 - total_weight, 2)
        max_pair = max(weight_pct, key=lambda k: weight_pct[k])
        weight_pct[max_pair] = round(weight_pct[max_pair] + diff, 2)

    # --- Fill rows ---
    start_row = 12
    total_rem = total_inf = total_app = 0
    actuals = []
    for idx in range(10):
        r = start_row + idx
        ws[f"A{r}"] = idx + 1
        if idx < len(subtopic_pairs):
            topic, sub = subtopic_pairs[idx]
            ws[f"B{r}"] = sub
            ws[f"C{r}"] = ""  # Competencies
            ws[f"D{r}"] = ""  # Time Spent / Frequency (min) - BLANK
            ws[f"E{r}"] = f"{weight_pct.get((topic, sub), 0)}%"

            n_rem = len(subtopic_map[(topic, sub)]["remembering"])
            n_inf = len(subtopic_map[(topic, sub)]["inferential"])
            n_app = len(subtopic_map[(topic, sub)]["applied"])
            total = n_rem + n_inf + n_app

            ws[f"F{r}"] = n_rem
            ws[f"G{r}"] = ", ".join(str(i) for i in subtopic_map[(topic, sub)]["remembering"])
            ws[f"H{r}"] = n_inf
            ws[f"I{r}"] = ", ".join(str(i) for i in subtopic_map[(topic, sub)]["inferential"])
            ws[f"J{r}"] = n_app
            ws[f"K{r}"] = ", ".join(str(i) for i in subtopic_map[(topic, sub)]["applied"])

            # --- Actual and Adjusted formulas ---
            ws[f"L{r}"] = f"=ROUND(VALUE(LEFT(E{r},LEN(E{r})-1))/100*50,2)"
            ws[f"M{r}"] = f"=ROUND(L{r},0)"
        else:
            for col in range(2, 15):
                ws[f"{get_column_letter(col)}{r}"] = ""
        for col in range(1, 15):
            ws[f"{get_column_letter(col)}{r}"].border = border
            ws[f"{get_column_letter(col)}{r}"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # --- Compute POI totals for each type ---
    total_poi_rem = sum(
        len(subtopic_map[pair]["remembering"]) for pair in subtopic_pairs
    )
    total_poi_inf = sum(
        len(subtopic_map[pair]["inferential"]) for pair in subtopic_pairs
    )
    total_poi_app = sum(
        len(subtopic_map[pair]["applied"]) for pair in subtopic_pairs
    )
    total_noi_rem = sum(
        len(subtopic_map[pair]["remembering"]) for pair in subtopic_pairs
    )
    total_noi_inf = sum(
        len(subtopic_map[pair]["inferential"]) for pair in subtopic_pairs
    )
    total_noi_app = sum(
        len(subtopic_map[pair]["applied"]) for pair in subtopic_pairs
    )

    # --- TOTAL ROW (after last subtopic row) ---
    total_row = start_row + 10  # Always after 10 rows
    ws[f"B{total_row}"] = "TOTAL"
    ws[f"D{total_row}"] = ""  # Time Spent blank
    ws[f"E{total_row}"] = "100%"  # Always show 100% for total weight

    # Leave NOI columns blank in total row
    ws[f"F{total_row}"] = ""
    ws[f"H{total_row}"] = ""
    ws[f"J{total_row}"] = ""
    ws[f"L{total_row}"] = ""
    ws[f"M{total_row}"] = ""

    # POI columns: show total count of POIs
    ws[f"G{total_row}"] = total_poi_rem
    ws[f"I{total_row}"] = total_poi_inf
    ws[f"K{total_row}"] = total_poi_app

    # Style the total row for emphasis (bold, black)
    for col in range(2, 14):
        cell = ws[f"{get_column_letter(col)}{total_row}"]
        cell.font = Font(bold=True, color="000000")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    # --- Add 50/50 row below Actual/Adjusted (move to row 22) ---
   
    actual_adjusted_row = total_row + 1  # This is row 22 if start_row=12
    ws[f"L{actual_adjusted_row}"] = f"=SUM(L{start_row}:L{total_row-1})"
    ws[f"M{actual_adjusted_row}"] = f"=SUM(M{start_row}:L{total_row-1})"
    ws[f"L{actual_adjusted_row}"].font = Font(bold=True, color="000000")
    ws[f"M{actual_adjusted_row}"].font = Font(bold=True, color="000000")
    ws[f"L{actual_adjusted_row}"].alignment = Alignment(horizontal="center", vertical="center")
    ws[f"M{actual_adjusted_row}"].alignment = Alignment(horizontal="center", vertical="center")

    # --- L23 and M23 merged, display actual total (accurate to actual/adjusted) ---
    merge_row = total_row + 2  # This is row 23 if start_row=12
    ws.merge_cells(f"L{merge_row}:M{merge_row}")
    ws[f"L{merge_row}"] = f"=SUM(L{start_row}:L{total_row-1})"
    ws[f"L{merge_row}"].font = Font(bold=True, color="000000")
    ws[f"L{merge_row}"].alignment = Alignment(horizontal="center", vertical="center")

    # --- Signature Block (define sig_row absolutely after merge_row) ---
    sig_row = merge_row + 5

    # Prepared by (left)
    ws[f"B{sig_row}"] = "Prepared by:"
    ws.merge_cells(f"B{sig_row+1}:D{sig_row+1}")
    ws[f"B{sig_row+1}"].border = Border(bottom=Side(style='thin'))
    ws[f"B{sig_row+2}"] = "Faculty"

    # Checked by (right)
    ws[f"I{sig_row}"] = "Checked by:"
    ws.merge_cells(f"I{sig_row+1}:K{sig_row+1}")
    ws[f"I{sig_row+1}"].border = Border(bottom=Side(style='thin'))
    ws[f"I{sig_row+2}"] = "Department Head"

    # Approved by (centered)
    ws[f"E{sig_row+2}"] = "Approved by:"
    ws.merge_cells(f"E{sig_row+3}:G{sig_row+3}")
    ws[f"E{sig_row+3}"].border = Border(bottom=Side(style='thin'))
    ws[f"E{sig_row+4}"] = "Dean"

    #

    for col, row in [
        ("B", sig_row), ("B", sig_row+2),
        ("I", sig_row), ("I", sig_row+2),
        ("E", sig_row+2), ("E", sig_row+4)
    ]:
        ws[f"{col}{row}"].alignment = Alignment(horizontal="center", vertical="center")

    # Center align signature lines
    for cell in [f"B{sig_row+1}", f"I{sig_row+1}", f"E{sig_row+3}"]:
        ws[cell].alignment = Alignment(horizontal="center", vertical="center")

    # Optionally, clear N column if not needed
    for idx in range(10):
        ws[f"N{start_row + idx}"] = ""

    # --- Save the workbook ---
    wb.save(output_path)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyleSheet(GLOBAL_STYLE)  
    window = LoginWindow()
    window.showMaximized()
    sys.exit(app.exec())
